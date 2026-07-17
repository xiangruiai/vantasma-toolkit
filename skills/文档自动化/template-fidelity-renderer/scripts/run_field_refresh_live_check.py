#!/usr/bin/env python3
import argparse
import json
import subprocess
import sys
import tempfile
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception as exc:
    Document = None
    OxmlElement = None
    qn = None
    DOCX_IMPORT_ERROR = exc
else:
    DOCX_IMPORT_ERROR = None


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def write_json(path, data):
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    Path(path).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def append_field(paragraph, instruction, stale_result):
    begin = OxmlElement("w:r")
    begin_char = OxmlElement("w:fldChar")
    begin_char.set(qn("w:fldCharType"), "begin")
    begin.append(begin_char)
    paragraph._p.append(begin)

    instr_run = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = f" {instruction} "
    instr_run.append(instr_text)
    paragraph._p.append(instr_run)

    separate = OxmlElement("w:r")
    separate_char = OxmlElement("w:fldChar")
    separate_char.set(qn("w:fldCharType"), "separate")
    separate.append(separate_char)
    paragraph._p.append(separate)

    result = paragraph.add_run(stale_result)

    end = OxmlElement("w:r")
    end_char = OxmlElement("w:fldChar")
    end_char.set(qn("w:fldCharType"), "end")
    end.append(end_char)
    paragraph._p.append(end)
    return result


def create_page_field_fixture(path):
    if Document is None:
        raise RuntimeError(f"python-docx is required to create the fixture: {DOCX_IMPORT_ERROR}")
    doc = Document()
    paragraph = doc.add_paragraph("Page field result: ")
    append_field(paragraph, "PAGE", "999")
    doc.add_page_break()
    doc.add_paragraph("Second page for page-field recalculation.")
    doc.save(path)


def docx_text(path):
    texts = []
    with ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    for node in root.findall(".//w:t", NS):
        texts.append(node.text or "")
    return "".join(texts)


def run_refresh(skill_scripts, source, output, report, backend, timeout, pdf_output=None):
    command = [
        sys.executable,
        str(skill_scripts / "refresh_fields.py"),
        str(source),
        "--output",
        str(output),
        "--report",
        str(report),
        "--backend",
        backend,
        "--timeout",
        str(timeout),
    ]
    if pdf_output:
        command.extend(["--pdf-output", str(pdf_output)])
    cp = subprocess.run(command, capture_output=True, text=True, timeout=timeout + 30)
    return {
        "command": command,
        "returncode": cp.returncode,
        "stdout": cp.stdout[-6000:],
        "stderr": cp.stderr[-6000:],
        "ok": cp.returncode == 0,
    }


def main():
    parser = argparse.ArgumentParser(description="Live-check refresh_fields.py against a DOCX with a stale PAGE field result.")
    parser.add_argument("--workdir", help="Directory for generated fixture and reports. Defaults to a temp directory.")
    parser.add_argument("--backend", choices=["auto", "request-only", "libreoffice-cli", "libreoffice-uno", "word-applescript"], default="auto")
    parser.add_argument("--report", help="Summary report path. Defaults to <workdir>/live-check.json.")
    parser.add_argument("--pdf", action="store_true", help="Also ask refresh_fields.py to export a sidecar PDF when the backend supports it.")
    parser.add_argument("--timeout", type=int, default=180)
    args = parser.parse_args()

    scripts = Path(__file__).resolve().parent
    if args.workdir:
        workdir = Path(args.workdir).expanduser().resolve()
        workdir.mkdir(parents=True, exist_ok=True)
    else:
        workdir = Path(tempfile.mkdtemp(prefix="template-fidelity-field-refresh-live-"))
    source = workdir / "stale-page-field.docx"
    output = workdir / "refreshed-page-field.docx"
    refresh_report = workdir / "refresh-fields.json"
    pdf_output = workdir / "refreshed-page-field.pdf" if args.pdf else None
    summary_report = Path(args.report).expanduser() if args.report else workdir / "live-check.json"

    summary = {
        "status": "unknown",
        "backend": args.backend,
        "workdir": str(workdir),
        "source": str(source),
        "output": str(output),
        "refresh_report": str(refresh_report),
        "pdf_output": str(pdf_output) if pdf_output else None,
        "source_text": None,
        "output_text": None,
        "stale_result": "999",
        "field_result_changed": False,
        "actual_field_results_refresh": False,
        "errors": [],
    }
    try:
        create_page_field_fixture(source)
        summary["source_text"] = docx_text(source)
        step = run_refresh(scripts, source, output, refresh_report, args.backend, args.timeout, pdf_output=pdf_output)
        summary["refresh_command"] = step
        if refresh_report.exists():
            refresh = json.loads(refresh_report.read_text(encoding="utf-8"))
        else:
            refresh = {}
        summary["refresh"] = refresh
        summary["actual_field_results_refresh"] = bool(refresh.get("actual_field_results_refresh"))
        if output.exists():
            summary["output_text"] = docx_text(output)
            summary["field_result_changed"] = "999" not in summary["output_text"] and summary["output_text"] != summary["source_text"]
        if args.backend == "request-only":
            summary["status"] = "passed" if step["ok"] and not summary["actual_field_results_refresh"] and not summary["field_result_changed"] else "failed"
        elif summary["actual_field_results_refresh"] and summary["field_result_changed"]:
            summary["status"] = "passed"
        elif args.backend == "auto" and refresh.get("selected_backend") == "request-only":
            summary["status"] = "warning"
            summary["errors"].append("No real field-refresh backend was available; auto fell back to request-only.")
        else:
            summary["status"] = "failed"
            summary["errors"].append("A real backend did not prove that the stale PAGE field result changed.")
    except Exception as exc:
        summary["status"] = "failed"
        summary["errors"].append(str(exc))
    finally:
        write_json(summary_report, summary)
        print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0 if summary["status"] in {"passed", "warning"} else 1


if __name__ == "__main__":
    raise SystemExit(main())
