#!/usr/bin/env python3
import argparse
import importlib.util
import json
import struct
import subprocess
import sys
import tempfile
import zipfile
import zlib
from pathlib import Path
from xml.dom import minidom

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn
    from docx.shared import Inches
except Exception as exc:
    Document = None
    WD_SECTION = None
    WD_STYLE_TYPE = None
    OxmlElement = None
    parse_xml = None
    qn = None
    Inches = None
    DOCX_IMPORT_ERROR = exc
else:
    DOCX_IMPORT_ERROR = None


SKILL_DIR = Path(__file__).resolve().parents[1]
SCRIPTS = SKILL_DIR / "scripts"

NOTE_TEST_CONFIGS = {
    "footnotes": {
        "part": "word/footnotes.xml",
        "root": "w:footnotes",
        "item": "w:footnote",
        "reference": "w:footnoteReference",
        "relationship_type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
    },
    "endnotes": {
        "part": "word/endnotes.xml",
        "root": "w:endnotes",
        "item": "w:endnote",
        "reference": "w:endnoteReference",
        "relationship_type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes",
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml",
    },
}

COMMENT_TEST_CONFIG = {
    "part": "word/comments.xml",
    "root": "w:comments",
    "item": "w:comment",
    "relationship_type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
    "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
}

CUSTOM_XML_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml"
CUSTOM_XML_PROPS_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXmlProps"
CUSTOM_XML_PROPS_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.customXmlProperties+xml"
W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"


def run_json(command):
    cp = subprocess.run(command, capture_output=True, text=True, timeout=120)
    if cp.returncode != 0:
        raise AssertionError(
            "Command failed:\n"
            + " ".join(str(part) for part in command)
            + f"\nreturncode={cp.returncode}\nstdout={cp.stdout}\nstderr={cp.stderr}"
        )
    return json.loads(cp.stdout)


def run_allow_warning(command):
    cp = subprocess.run(command, capture_output=True, text=True, timeout=120)
    if cp.returncode not in {0, 1}:
        raise AssertionError(
            "Command failed unexpectedly:\n"
            + " ".join(str(part) for part in command)
            + f"\nreturncode={cp.returncode}\nstdout={cp.stdout}\nstderr={cp.stderr}"
        )
    return json.loads(cp.stdout)


def load_script_module(name):
    module_path = SCRIPTS / name
    spec = importlib.util.spec_from_file_location(module_path.stem, module_path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def check(condition, message):
    if not condition:
        raise AssertionError(message)


def require_python_docx():
    if Document is None:
        raise AssertionError(f"python-docx is required for fixture generation: {DOCX_IMPORT_ERROR}")


def set_split_runs(paragraph, chunks):
    paragraph.text = ""
    for index, chunk in enumerate(chunks):
        run = paragraph.add_run(chunk)
        if index % 2:
            run.bold = True


def write_png(path, rgb, width=8, height=8):
    def chunk(name, data):
        body = name + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    rows = b"".join(b"\x00" + bytes(rgb) * width for _ in range(height))
    raw = b"\x89PNG\r\n\x1a\n"
    raw += chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
    raw += chunk(b"IDAT", zlib.compress(rows))
    raw += chunk(b"IEND", b"")
    path.write_bytes(raw)


def write_gif(path):
    path.write_bytes(
        b"GIF89a\x01\x00\x01\x00\x80\x00\x00\x00\x00\x00\xff\xff\xff"
        b"!\xf9\x04\x01\x00\x00\x00\x00,\x00\x00\x00\x00\x01\x00\x01"
        b"\x00\x00\x02\x02D\x01\x00;"
    )


def add_numbering_definition(doc, num_id=42, abstract_id=42, lvl_text_value="Ref-%1."):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), lvl_text_value)
    lvl.append(start)
    lvl.append(num_fmt)
    lvl.append(lvl_text)
    abstract.append(lvl)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(abstract)
    numbering.append(num)


def set_style_numbering(style, num_id, ilvl_value="0"):
    style_element = style.element
    p_pr = style_element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        style_element.append(p_pr)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), ilvl_value)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def add_external_hyperlink_paragraph(doc, text, url):
    paragraph = doc.add_paragraph()
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return paragraph


def xml_text(node):
    return "".join(child.firstChild.data for child in node.getElementsByTagName("w:t") if child.firstChild)


def next_rid(rels_dom):
    values = []
    for node in rels_dom.getElementsByTagName("Relationship"):
        rid = node.getAttribute("Id")
        if rid.startswith("rId") and rid[3:].isdigit():
            values.append(int(rid[3:]))
    return f"rId{max(values, default=0) + 1}"


def patch_docx_with_note(path, paragraph_text, note_text, note_id, kind):
    config = NOTE_TEST_CONFIGS[kind]
    with zipfile.ZipFile(path) as archive:
        package = {name: archive.read(name) for name in archive.namelist()}

    document_dom = minidom.parseString(package["word/document.xml"])
    target_paragraph = None
    for paragraph in document_dom.getElementsByTagName("w:p"):
        if paragraph_text in xml_text(paragraph):
            target_paragraph = paragraph
            break
    check(target_paragraph is not None, f"Could not find paragraph for {kind} fixture.")
    run = document_dom.createElement("w:r")
    reference = document_dom.createElement(config["reference"])
    reference.setAttribute("w:id", str(note_id))
    run.appendChild(reference)
    target_paragraph.appendChild(run)
    package["word/document.xml"] = document_dom.toxml(encoding="utf-8")

    notes_dom = minidom.Document()
    root = notes_dom.createElement(config["root"])
    root.setAttribute("xmlns:w", "http://schemas.openxmlformats.org/wordprocessingml/2006/main")
    notes_dom.appendChild(root)
    note = notes_dom.createElement(config["item"])
    note.setAttribute("w:id", str(note_id))
    paragraph = notes_dom.createElement("w:p")
    note_run = notes_dom.createElement("w:r")
    text_node = notes_dom.createElement("w:t")
    text_node.appendChild(notes_dom.createTextNode(note_text))
    note_run.appendChild(text_node)
    paragraph.appendChild(note_run)
    note.appendChild(paragraph)
    root.appendChild(note)
    package[config["part"]] = notes_dom.toxml(encoding="utf-8")

    content_types_dom = minidom.parseString(package["[Content_Types].xml"])
    part_name = "/" + config["part"]
    if not any(node.getAttribute("PartName") == part_name for node in content_types_dom.getElementsByTagName("Override")):
        override = content_types_dom.createElement("Override")
        override.setAttribute("PartName", part_name)
        override.setAttribute("ContentType", config["content_type"])
        content_types_dom.documentElement.appendChild(override)
    package["[Content_Types].xml"] = content_types_dom.toxml(encoding="utf-8")

    rels_dom = minidom.parseString(package["word/_rels/document.xml.rels"])
    if not any(node.getAttribute("Type") == config["relationship_type"] for node in rels_dom.getElementsByTagName("Relationship")):
        relationship = rels_dom.createElement("Relationship")
        relationship.setAttribute("Id", next_rid(rels_dom))
        relationship.setAttribute("Type", config["relationship_type"])
        relationship.setAttribute("Target", Path(config["part"]).name)
        rels_dom.documentElement.appendChild(relationship)
    package["word/_rels/document.xml.rels"] = rels_dom.toxml(encoding="utf-8")

    tmp_path = path.with_suffix(".patched.docx")
    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, content in package.items():
            archive.writestr(name, content)
    tmp_path.replace(path)


def ensure_content_type_override(content_types_dom, part, content_type):
    part_name = "/" + part
    for node in content_types_dom.getElementsByTagName("Override"):
        if node.getAttribute("PartName") == part_name:
            node.setAttribute("ContentType", content_type)
            return
    override = content_types_dom.createElement("Override")
    override.setAttribute("PartName", part_name)
    override.setAttribute("ContentType", content_type)
    content_types_dom.documentElement.appendChild(override)


def patch_docx_with_custom_xml(path, store_item_id, item_xml):
    with zipfile.ZipFile(path) as archive:
        package = {name: archive.read(name) for name in archive.namelist()}

    package["customXml/item1.xml"] = item_xml.encode("utf-8")
    props_dom = minidom.Document()
    root = props_dom.createElement("ds:datastoreItem")
    root.setAttribute("ds:itemID", store_item_id)
    root.setAttribute("xmlns:ds", "http://schemas.openxmlformats.org/officeDocument/2006/customXml")
    schema_refs = props_dom.createElement("ds:schemaRefs")
    root.appendChild(schema_refs)
    props_dom.appendChild(root)
    package["customXml/itemProps1.xml"] = props_dom.toxml(encoding="utf-8")

    item_rels_dom = minidom.Document()
    rels_root = item_rels_dom.createElement("Relationships")
    rels_root.setAttribute("xmlns", "http://schemas.openxmlformats.org/package/2006/relationships")
    item_rels_dom.appendChild(rels_root)
    relationship = item_rels_dom.createElement("Relationship")
    relationship.setAttribute("Id", "rId1")
    relationship.setAttribute("Type", CUSTOM_XML_PROPS_REL_TYPE)
    relationship.setAttribute("Target", "itemProps1.xml")
    rels_root.appendChild(relationship)
    package["customXml/_rels/item1.xml.rels"] = item_rels_dom.toxml(encoding="utf-8")

    content_types_dom = minidom.parseString(package["[Content_Types].xml"])
    ensure_content_type_override(content_types_dom, "customXml/item1.xml", "application/xml")
    ensure_content_type_override(content_types_dom, "customXml/itemProps1.xml", CUSTOM_XML_PROPS_CONTENT_TYPE)
    package["[Content_Types].xml"] = content_types_dom.toxml(encoding="utf-8")

    package_rels_dom = minidom.parseString(package["_rels/.rels"])
    if not any(node.getAttribute("Type") == CUSTOM_XML_REL_TYPE and node.getAttribute("Target") == "customXml/item1.xml" for node in package_rels_dom.getElementsByTagName("Relationship")):
        relationship = package_rels_dom.createElement("Relationship")
        relationship.setAttribute("Id", next_rid(package_rels_dom))
        relationship.setAttribute("Type", CUSTOM_XML_REL_TYPE)
        relationship.setAttribute("Target", "customXml/item1.xml")
        package_rels_dom.documentElement.appendChild(relationship)
    package["_rels/.rels"] = package_rels_dom.toxml(encoding="utf-8")

    tmp_path = path.with_suffix(".patched.docx")
    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, content in package.items():
            archive.writestr(name, content)
    tmp_path.replace(path)


def patch_docx_with_comment(path, paragraph_text, comment_text, comment_id):
    with zipfile.ZipFile(path) as archive:
        package = {name: archive.read(name) for name in archive.namelist()}

    document_dom = minidom.parseString(package["word/document.xml"])
    target_paragraph = None
    for paragraph in document_dom.getElementsByTagName("w:p"):
        if paragraph_text in xml_text(paragraph):
            target_paragraph = paragraph
            break
    check(target_paragraph is not None, "Could not find paragraph for comment fixture.")
    range_start = document_dom.createElement("w:commentRangeStart")
    range_start.setAttribute("w:id", str(comment_id))
    range_end = document_dom.createElement("w:commentRangeEnd")
    range_end.setAttribute("w:id", str(comment_id))
    reference_run = document_dom.createElement("w:r")
    comment_reference = document_dom.createElement("w:commentReference")
    comment_reference.setAttribute("w:id", str(comment_id))
    reference_run.appendChild(comment_reference)
    first_child = target_paragraph.firstChild
    if first_child is not None:
        target_paragraph.insertBefore(range_start, first_child)
    else:
        target_paragraph.appendChild(range_start)
    target_paragraph.appendChild(range_end)
    target_paragraph.appendChild(reference_run)
    package["word/document.xml"] = document_dom.toxml(encoding="utf-8")

    comments_dom = minidom.Document()
    root = comments_dom.createElement(COMMENT_TEST_CONFIG["root"])
    root.setAttribute("xmlns:w", "http://schemas.openxmlformats.org/wordprocessingml/2006/main")
    comments_dom.appendChild(root)
    comment = comments_dom.createElement(COMMENT_TEST_CONFIG["item"])
    comment.setAttribute("w:id", str(comment_id))
    comment.setAttribute("w:author", "Reference Reviewer")
    comment.setAttribute("w:date", "2026-07-16T00:00:00Z")
    paragraph = comments_dom.createElement("w:p")
    run = comments_dom.createElement("w:r")
    text_node = comments_dom.createElement("w:t")
    text_node.appendChild(comments_dom.createTextNode(comment_text))
    run.appendChild(text_node)
    paragraph.appendChild(run)
    comment.appendChild(paragraph)
    root.appendChild(comment)
    package[COMMENT_TEST_CONFIG["part"]] = comments_dom.toxml(encoding="utf-8")

    content_types_dom = minidom.parseString(package["[Content_Types].xml"])
    part_name = "/" + COMMENT_TEST_CONFIG["part"]
    if not any(node.getAttribute("PartName") == part_name for node in content_types_dom.getElementsByTagName("Override")):
        override = content_types_dom.createElement("Override")
        override.setAttribute("PartName", part_name)
        override.setAttribute("ContentType", COMMENT_TEST_CONFIG["content_type"])
        content_types_dom.documentElement.appendChild(override)
    package["[Content_Types].xml"] = content_types_dom.toxml(encoding="utf-8")

    rels_dom = minidom.parseString(package["word/_rels/document.xml.rels"])
    if not any(node.getAttribute("Type") == COMMENT_TEST_CONFIG["relationship_type"] for node in rels_dom.getElementsByTagName("Relationship")):
        relationship = rels_dom.createElement("Relationship")
        relationship.setAttribute("Id", next_rid(rels_dom))
        relationship.setAttribute("Type", COMMENT_TEST_CONFIG["relationship_type"])
        relationship.setAttribute("Target", Path(COMMENT_TEST_CONFIG["part"]).name)
        rels_dom.documentElement.appendChild(relationship)
    package["word/_rels/document.xml.rels"] = rels_dom.toxml(encoding="utf-8")

    tmp_path = path.with_suffix(".patched.docx")
    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, content in package.items():
            archive.writestr(name, content)
    tmp_path.replace(path)


def patch_docx_with_complex_reference_structures(path, paragraph_text):
    with zipfile.ZipFile(path) as archive:
        package = {name: archive.read(name) for name in archive.namelist()}

    document_dom = minidom.parseString(package["word/document.xml"])
    target_paragraph = None
    for paragraph in document_dom.getElementsByTagName("w:p"):
        if paragraph_text in xml_text(paragraph):
            target_paragraph = paragraph
            break
    check(target_paragraph is not None, "Could not find paragraph for complex reference fixture.")

    field = document_dom.createElement("w:fldSimple")
    field.setAttribute("w:instr", "PAGE")
    field_run = document_dom.createElement("w:r")
    field_text = document_dom.createElement("w:t")
    field_text.appendChild(document_dom.createTextNode("1"))
    field_run.appendChild(field_text)
    field.appendChild(field_run)
    target_paragraph.appendChild(field)

    insertion = document_dom.createElement("w:ins")
    insertion.setAttribute("w:id", "21")
    insertion.setAttribute("w:author", "Reference Reviewer")
    insertion.setAttribute("w:date", "2026-07-16T00:00:00Z")
    insertion_run = document_dom.createElement("w:r")
    insertion_text = document_dom.createElement("w:t")
    insertion_text.appendChild(document_dom.createTextNode("Inserted revision"))
    insertion_run.appendChild(insertion_text)
    insertion.appendChild(insertion_run)
    target_paragraph.appendChild(insertion)

    package["word/document.xml"] = document_dom.toxml(encoding="utf-8")

    tmp_path = path.with_suffix(".patched.docx")
    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, content in package.items():
            archive.writestr(name, content)
    tmp_path.replace(path)


def add_custom_numbered_paragraph(doc, text, num_id=42, abstract_id=42):
    add_numbering_definition(doc, num_id=num_id, abstract_id=abstract_id)
    paragraph = doc.add_paragraph(text)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    return paragraph


def content_control_pr(tag=None, alias=None, binding=None, lock=None, text_control=None):
    sdt_pr = OxmlElement("w:sdtPr")
    if alias:
        alias_node = OxmlElement("w:alias")
        alias_node.set(qn("w:val"), alias)
        sdt_pr.append(alias_node)
    if tag:
        tag_node = OxmlElement("w:tag")
        tag_node.set(qn("w:val"), tag)
        sdt_pr.append(tag_node)
    if binding:
        binding_node = OxmlElement("w:dataBinding")
        if binding.get("prefix_mappings"):
            binding_node.set(qn("w:prefixMappings"), binding["prefix_mappings"])
        if binding.get("xpath"):
            binding_node.set(qn("w:xpath"), binding["xpath"])
        if binding.get("store_item_id"):
            binding_node.set(qn("w:storeItemID"), binding["store_item_id"])
        sdt_pr.append(binding_node)
    if lock:
        lock_node = OxmlElement("w:lock")
        lock_node.set(qn("w:val"), lock)
        sdt_pr.append(lock_node)
    if text_control:
        text_node = OxmlElement("w:text")
        if "multi_line" in text_control:
            text_node.set(qn("w:multiLine"), "1" if text_control["multi_line"] else "0")
        sdt_pr.append(text_node)
    return sdt_pr


def wrap_paragraph_in_content_control(paragraph, tag=None, alias=None, text_control=None):
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    element = paragraph._p
    parent = element.getparent()
    index = parent.index(element)
    parent.remove(element)
    content.append(element)
    sdt.append(content_control_pr(tag=tag, alias=alias, text_control=text_control))
    sdt.append(content)
    parent.insert(index, sdt)
    return sdt


def mark_content_control_showing_placeholder(sdt):
    sdt_pr = next((child for child in sdt if child.tag == qn("w:sdtPr")), None)
    check(sdt_pr is not None, "Content control is missing sdtPr.")
    sdt_pr.append(OxmlElement("w:showingPlcHdr"))


def add_run_content_control(paragraph, text, tag=None, alias=None, lock=None):
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    content.append(run)
    sdt.append(content_control_pr(tag=tag, alias=alias, lock=lock))
    sdt.append(content)
    paragraph._p.append(sdt)
    return sdt


def build_styled_run(text="", ascii_font="Courier New", east_asia_font="SimSun", size_val="28", bold=False):
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    fonts.set(qn("w:eastAsia"), east_asia_font)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), size_val)
    run_pr.append(fonts)
    run_pr.append(size)
    if bold:
        run_pr.append(OxmlElement("w:b"))
    run.append(run_pr)
    if text:
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(text_node)
    return run


def add_styled_run(paragraph, text="", ascii_font="Courier New", east_asia_font="SimSun", size_val="28", bold=False):
    run = build_styled_run(text, ascii_font=ascii_font, east_asia_font=east_asia_font, size_val=size_val, bold=bold)
    paragraph._p.append(run)
    return run


def add_empty_styled_run_content_control(paragraph, tag=None, alias=None):
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    run = build_styled_run(bold=True)
    content.append(run)
    sdt.append(content_control_pr(tag=tag, alias=alias))
    sdt.append(content)
    paragraph._p.append(sdt)
    return sdt


def add_bookmark_text(paragraph, name, text="", bookmark_id=1):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    paragraph._p.append(start)
    if text:
        run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(text_node)
        paragraph._p.append(run)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)
    return start, end


def add_field_code(paragraph, instruction="TOC \\o \"1-3\" \\h \\z \\u"):
    run_begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin.append(fld_begin)
    paragraph._p.append(run_begin)

    run_instruction = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.text = instruction
    run_instruction.append(instr)
    paragraph._p.append(run_instruction)

    run_separate = OxmlElement("w:r")
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run_separate.append(fld_separate)
    paragraph._p.append(run_separate)

    paragraph.add_run("Generated field result")

    run_end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end.append(fld_end)
    paragraph._p.append(run_end)


def add_office_math(paragraph):
    if parse_xml is None:
        raise AssertionError("parse_xml is required for math fixture generation.")
    paragraph._p.append(
        parse_xml(
            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
            "<m:r><m:t>x+y</m:t></m:r>"
            "</m:oMath>"
        )
    )


def add_databound_run_content_control(paragraph, text, xpath, store_item_id, prefix_mappings=None):
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    content.append(run)
    sdt.append(
        content_control_pr(
            binding={
                "xpath": xpath,
                "store_item_id": store_item_id,
                "prefix_mappings": prefix_mappings,
            }
        )
    )
    sdt.append(content)
    paragraph._p.append(sdt)
    return sdt


def add_checkbox_content_control(paragraph, tag=None, alias=None, checked=False):
    sdt = OxmlElement("w:sdt")
    sdt_pr = content_control_pr(tag=tag, alias=alias)
    checkbox = OxmlElement("w14:checkbox")
    checked_el = OxmlElement("w14:checked")
    checked_el.set(qn("w14:val"), "1" if checked else "0")
    checked_state = OxmlElement("w14:checkedState")
    checked_state.set(qn("w14:val"), "2612")
    checked_state.set(qn("w14:font"), "MS Gothic")
    unchecked_state = OxmlElement("w14:uncheckedState")
    unchecked_state.set(qn("w14:val"), "2610")
    unchecked_state.set(qn("w14:font"), "MS Gothic")
    checkbox.append(checked_el)
    checkbox.append(checked_state)
    checkbox.append(unchecked_state)
    sdt_pr.append(checkbox)
    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = "☒" if checked else "☐"
    run.append(text_node)
    content.append(run)
    sdt.append(sdt_pr)
    sdt.append(content)
    paragraph._p.append(sdt)
    return sdt


def add_choice_content_control(paragraph, tag=None, alias=None, kind="dropDownList", options=None, selected=None):
    options = options or []
    sdt = OxmlElement("w:sdt")
    sdt_pr = content_control_pr(tag=tag, alias=alias)
    choice = OxmlElement(f"w:{kind}")
    selected_text = selected
    for display_text, value in options:
        item = OxmlElement("w:listItem")
        item.set(qn("w:displayText"), display_text)
        item.set(qn("w:value"), value)
        choice.append(item)
        if selected_text is None:
            selected_text = display_text
    sdt_pr.append(choice)
    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = selected_text or ""
    run.append(text_node)
    content.append(run)
    sdt.append(sdt_pr)
    sdt.append(content)
    paragraph._p.append(sdt)
    return sdt


def add_date_content_control(paragraph, tag=None, alias=None, selected="2026-01-01", full_date="2026-01-01T00:00:00Z", date_format="yyyy-MM-dd"):
    sdt = OxmlElement("w:sdt")
    sdt_pr = content_control_pr(tag=tag, alias=alias)
    date_node = OxmlElement("w:date")
    date_node.set(qn("w:fullDate"), full_date)
    date_format_node = OxmlElement("w:dateFormat")
    date_format_node.set(qn("w:val"), date_format)
    lid_node = OxmlElement("w:lid")
    lid_node.set(qn("w:val"), "zh-CN")
    store_node = OxmlElement("w:storeMappedDataAs")
    store_node.set(qn("w:val"), "dateTime")
    calendar_node = OxmlElement("w:calendar")
    calendar_node.set(qn("w:val"), "gregorian")
    date_node.append(date_format_node)
    date_node.append(lid_node)
    date_node.append(store_node)
    date_node.append(calendar_node)
    sdt_pr.append(date_node)
    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = selected
    run.append(text_node)
    content.append(run)
    sdt.append(sdt_pr)
    sdt.append(content)
    paragraph._p.append(sdt)
    return sdt


def add_repeating_section_content_control(doc, tag=None, alias=None):
    sdt = OxmlElement("w:sdt")
    sdt_pr = content_control_pr(tag=tag, alias=alias)
    sdt_pr.append(parse_xml(f'<w15:repeatingSection xmlns:w15="{W15_NS}"/>'))
    content = OxmlElement("w:sdtContent")
    item = OxmlElement("w:sdt")
    item_pr = content_control_pr()
    item_pr.append(parse_xml(f'<w15:repeatingSectionItem xmlns:w15="{W15_NS}"/>'))
    item_content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = "{{name}} x {{qty}}"
    run.append(text_node)
    paragraph.append(run)
    item_content.append(paragraph)
    item.append(item_pr)
    item.append(item_content)
    content.append(item)
    sdt.append(sdt_pr)
    sdt.append(content)
    doc._body._element.append(sdt)
    return sdt


def docx_xml_text(path, part):
    with zipfile.ZipFile(path) as archive:
        return archive.read(part).decode("utf-8", errors="ignore")


def docx_media_bytes(path):
    with zipfile.ZipFile(path) as archive:
        return {name: archive.read(name) for name in archive.namelist() if name.startswith("word/media/")}


def check_named(report, name):
    for row in report.get("checks", []):
        if row.get("name") == name:
            return row
    raise AssertionError(f"Missing check: {name}")


def test_header_footer_text_structure(workdir):
    require_python_docx()
    template = workdir / "header-template.docx"
    data = workdir / "header-data.json"
    output = workdir / "header-output.docx"
    render_report = workdir / "header-render.json"
    verify_report = workdir / "header-verify.json"

    doc = Document()
    section = doc.sections[0]
    set_split_runs(section.header.paragraphs[0], ["Client: {{client", "_name}}"])
    section.footer.paragraphs[0].text = "Confidential"
    set_split_runs(doc.add_paragraph(), ["Project: {{project", "_name}}"])
    doc.save(template)

    data.write_text(
        json.dumps({"fields": {"client_name": "Acme Corp", "project_name": "Template fidelity"}}, indent=2),
        encoding="utf-8",
    )
    run_json(
        [
            sys.executable,
            str(SCRIPTS / "render_docx.py"),
            str(template),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(render_report),
            "--strict",
        ]
    )
    report = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(verify_report),
            "--render-report",
            str(render_report),
        ]
    )

    check("Acme Corp" in docx_xml_text(output, "word/header1.xml"), "Header placeholder was not replaced.")
    check(check_named(report, "unresolved_placeholders")["passed"], "Rendered DOCX still has unresolved placeholders.")
    check(check_named(report, "invariant_parts_unchanged")["passed"], "Core invariant OOXML parts changed.")
    check(check_named(report, "text_part_structure_preserved")["passed"], "Header/footer structure changed.")
    check(check_named(report, "tracked_text_format_preserved")["passed"], "Tracked L1 text format was not preserved.")
    check(report["tracked_text_format"]["checked_count"] >= 2, "L1 render report did not expose tracked format checks.")
    text_parts = [row for row in report.get("part_comparison", []) if row.get("comparison") == "text_structure"]
    check(any(row.get("part") == "word/header1.xml" and row.get("text_changed") for row in text_parts), "Header text change was not reported as a text-part change.")

    bad_render_report = workdir / "bad-format-render.json"
    bad_verify_report = workdir / "bad-format-verify.json"
    bad_render_report.write_text(
        json.dumps(
            {
                "text_format_checks": [
                    {
                        "part": "word/document.xml",
                        "scope": "part_text_format_structure",
                        "preserved": False,
                    }
                ]
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    bad = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(bad_verify_report),
            "--render-report",
            str(bad_render_report),
        ]
    )
    check(not check_named(bad, "tracked_text_format_preserved")["passed"], "Verifier did not fail a tracked format drift.")


def test_cross_run_placeholders(workdir):
    require_python_docx()
    template = workdir / "cross-run-template.docx"
    data = workdir / "cross-run-data.json"
    output = workdir / "cross-run-output.docx"
    render_report = workdir / "cross-run-render.json"
    spec = workdir / "cross-run-spec.json"
    spec_output = workdir / "cross-run-spec-output.docx"
    spec_report = workdir / "cross-run-spec-render.json"

    doc = Document()
    set_split_runs(doc.add_paragraph(), ["Client: {{client", "_name}}"])
    set_split_runs(doc.add_paragraph(), ["Proposal: {{proposal", "_title}}"])
    doc.save(template)

    data.write_text(
        json.dumps({"fields": {"client_name": "Acme Corp", "proposal_title": "Automation Plan"}}, indent=2),
        encoding="utf-8",
    )
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "render_docx.py"),
            str(template),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(render_report),
            "--strict",
        ]
    )
    check(report["replacement_counts"]["{{client_name}}"] == 1, "render_docx did not count a cross-run client token.")
    check(report["replacement_counts"]["{{proposal_title}}"] == 1, "render_docx did not count a cross-run proposal token.")
    document_xml = docx_xml_text(output, "word/document.xml")
    check("Acme Corp" in document_xml and "Automation Plan" in document_xml, "render_docx did not replace cross-run tokens.")
    check("{{client" not in document_xml and "_name}}" not in document_xml, "render_docx left split token fragments.")

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "proposal_title",
                        "locator_type": "placeholder",
                        "locator": {"token": "{{proposal_title}}"},
                        "replacement_mode": "token",
                        "required": True,
                    }
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    spec_report_json = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(spec_output),
            "--report",
            str(spec_report),
        ]
    )
    field = spec_report_json["fields"][0]
    check(field["status"] == "filled", "fill_by_spec did not fill a cross-run placeholder.")
    check(field["token_counts"]["{{proposal_title}}"] == 1, "fill_by_spec did not count the cross-run placeholder.")
    spec_xml = docx_xml_text(spec_output, "word/document.xml")
    check("Automation Plan" in spec_xml and "{{proposal" not in spec_xml, "fill_by_spec left split token fragments.")


def test_multiline_line_break_rendering(workdir):
    require_python_docx()
    template = workdir / "multiline-template.docx"
    data = workdir / "multiline-data.json"
    render_data = workdir / "multiline-render-data.json"
    output = workdir / "multiline-output.docx"
    render_report = workdir / "multiline-render.json"
    verify_report = workdir / "multiline-verify.json"
    spec = workdir / "multiline-spec.json"
    spec_output = workdir / "multiline-spec-output.docx"
    spec_report = workdir / "multiline-spec-render.json"
    spec_verify_report = workdir / "multiline-spec-verify.json"

    doc = Document()
    doc.add_paragraph("Abstract: {{abstract}}")
    doc.add_paragraph("Clause placeholder")
    notes_paragraph = doc.add_paragraph("Notes: ")
    add_run_content_control(notes_paragraph, "Placeholder Notes", tag="notes", alias="Notes")
    doc.save(template)

    data.write_text(
        json.dumps(
            {
                "fields": {
                    "abstract": "First line\nSecond line",
                    "clause": "Clause A\nClause B",
                    "notes": "Note A\nNote B",
                },
                "literal_replacements": {"Clause placeholder": "Clause A\nClause B"},
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    render_data.write_text(
        json.dumps(
            {
                "fields": {"abstract": "First line\nSecond line"},
                "literal_replacements": {"Clause placeholder": "Clause A\nClause B"},
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "render_docx.py"),
            str(template),
            "--data",
            str(render_data),
            "--output",
            str(output),
            "--report",
            str(render_report),
            "--strict",
        ]
    )
    check(report["line_breaks_inserted"]["word/document.xml"] == 2, "render_docx did not report inserted line breaks.")
    document_xml = docx_xml_text(output, "word/document.xml")
    check(document_xml.count("<w:br") >= 2, "render_docx did not materialize multiline values as w:br.")
    check("First line" in document_xml and "Second line" in document_xml, "render_docx multiline field text is missing.")
    verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(verify_report),
            "--render-report",
            str(render_report),
        ]
    )
    check(check_named(verify, "tracked_text_format_preserved")["passed"], "Verifier treated expected render_docx line breaks as format drift.")

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "abstract",
                        "locator_type": "placeholder",
                        "locator": {"token": "{{abstract}}"},
                        "replacement_mode": "token",
                        "required": True,
                    },
                    {
                        "key": "clause",
                        "locator_type": "literal",
                        "locator": {"text": "Clause placeholder"},
                        "required": True,
                    },
                    {
                        "key": "notes",
                        "locator_type": "content_control",
                        "locator": {"tag": "notes"},
                        "required": True,
                    },
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    spec_render = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(spec_output),
            "--report",
            str(spec_report),
        ]
    )
    fields = {row["key"]: row for row in spec_render["fields"]}
    for key in ["abstract", "clause", "notes"]:
        check(fields[key]["status"] == "filled", f"fill_by_spec did not fill multiline field {key}.")
        check(fields[key].get("line_breaks_inserted") == 1, f"fill_by_spec did not report one line break for {key}.")
        check(fields[key]["format_check"].get("expected_structure_change") == "line_breaks_inserted", f"fill_by_spec did not mark expected line break structure for {key}.")
    spec_xml = docx_xml_text(spec_output, "word/document.xml")
    check(spec_xml.count("<w:br") >= 3, "fill_by_spec did not materialize multiline values as w:br.")
    check("Note A" in spec_xml and "Note B" in spec_xml, "fill_by_spec multiline content-control text is missing.")
    spec_verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(spec_output),
            "--report",
            str(spec_verify_report),
            "--render-report",
            str(spec_report),
        ]
    )
    check(check_named(spec_verify, "tracked_text_format_preserved")["passed"], "Verifier treated expected fill_by_spec line breaks as format drift.")
    check(check_named(spec_verify, "content_controls_preserved")["passed"], "Verifier did not preserve content control after multiline fill.")


def test_spec_bookmark_text_replace(workdir):
    require_python_docx()
    template = workdir / "bookmark-template.docx"
    analysis_report = workdir / "bookmark-analysis.json"
    draft_spec_output = workdir / "bookmark-draft-spec.json"
    draft_report = workdir / "bookmark-draft-report.json"
    spec = workdir / "bookmark-spec.json"
    data = workdir / "bookmark-data.json"
    output = workdir / "bookmark-output.docx"
    render_report = workdir / "bookmark-render.json"
    verify_report = workdir / "bookmark-verify.json"

    doc = Document()
    client_para = doc.add_paragraph()
    client_para.add_run("Client: ")
    add_bookmark_text(client_para, "client_name", "Old Client", bookmark_id=11)
    client_para.add_run(" / approved")
    code_para = doc.add_paragraph()
    code_para.add_run("Project code: ")
    add_styled_run(code_para, "", ascii_font="Courier New", east_asia_font="SimSun", size_val="28", bold=True)
    add_bookmark_text(code_para, "project_code", "", bookmark_id=12)
    code_para.add_run(" / end")
    doc.save(template)

    analysis = run_json(
        [
            sys.executable,
            str(SCRIPTS / "analyze_template.py"),
            str(template),
            "--output",
            str(analysis_report),
        ]
    )
    bookmarks = {row.get("name"): row for row in analysis.get("bookmarks", [])}
    check(bookmarks["client_name"]["spec_candidate"]["replacement_mode"] == "bookmark_text", "Analyze did not propose bookmark_text for named bookmark.")
    check(bookmarks["project_code"]["same_paragraph"], "Analyze did not report the empty bookmark as same-paragraph.")

    run_json(
        [
            sys.executable,
            str(SCRIPTS / "draft_spec.py"),
            str(template),
            "--spec-output",
            str(draft_spec_output),
            "--report",
            str(draft_report),
        ]
    )
    draft = json.loads(draft_spec_output.read_text(encoding="utf-8"))
    draft_fields = {row["key"]: row for row in draft["fields"]}
    check(draft_fields["client_name"]["replacement_mode"] == "bookmark_text", "Draft spec did not preserve bookmark_text mode.")
    check(draft_fields["project_code"]["locator_type"] == "bookmark", "Draft spec did not add the empty bookmark field.")

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "client_name",
                        "locator_type": "bookmark",
                        "locator": {"name": "client_name"},
                        "required": True,
                    },
                    {
                        "key": "project_code",
                        "locator_type": "bookmark",
                        "locator": {"name": "project_code"},
                        "required": True,
                    },
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(
        json.dumps({"fields": {"client_name": "New Client", "project_code": "PRJ-001"}}, indent=2),
        encoding="utf-8",
    )
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(render_report),
        ]
    )
    fields = {row["key"]: row for row in report["fields"]}
    check(fields["client_name"]["action"] == "bookmark_text_replace", "Bookmark text field used the wrong action.")
    check(fields["client_name"]["format_check"]["preserved"], "Bookmark text replacement did not preserve paragraph formatting.")
    check(fields["project_code"]["run_inserted"], "Empty bookmark did not report run insertion.")
    check(fields["project_code"]["run_style_source"] == "previous_run", "Empty bookmark did not inherit style from the nearest previous run.")
    check(fields["project_code"]["format_check"]["expected_structure_change"] == "bookmark_empty_range_inserted", "Empty bookmark insertion did not mark expected structure change.")
    document_xml = docx_xml_text(output, "word/document.xml")
    check("New Client" in document_xml and "Old Client" not in document_xml, "Bookmark text replacement failed.")
    check("Project code:" in document_xml and "PRJ-001" in document_xml and "/ end" in document_xml, "Empty bookmark insertion failed.")
    check('w:name="client_name"' in document_xml and 'w:name="project_code"' in document_xml, "Bookmark metadata was not preserved.")
    dom = minidom.parseString(document_xml.encode("utf-8"))
    project_text = next(
        (
            node
            for node in dom.getElementsByTagName("w:t")
            if node.firstChild and node.firstChild.data == "PRJ-001"
        ),
        None,
    )
    check(project_text is not None, "Could not find inserted bookmark text node.")
    project_run = project_text.parentNode
    project_run_pr = project_run.getElementsByTagName("w:rPr")
    check(project_run_pr, "Inserted bookmark run did not preserve a run property node.")
    project_run_xml = project_run.toxml()
    check("Courier New" in project_run_xml and "SimSun" in project_run_xml and 'w:val="28"' in project_run_xml and "<w:b" in project_run_xml, "Inserted bookmark run did not inherit neighboring font/size/bold style.")

    verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(verify_report),
            "--render-report",
            str(render_report),
        ]
    )
    check(check_named(verify, "tracked_text_format_preserved")["passed"], "Verifier treated expected bookmark insertion as format drift.")


def test_spec_hyperlink_field(workdir):
    require_python_docx()
    template = workdir / "hyperlink-template.docx"
    analysis_report = workdir / "hyperlink-analysis.json"
    draft_report = workdir / "hyperlink-draft-report.json"
    spec = workdir / "hyperlink-spec.json"
    data = workdir / "hyperlink-data.json"
    output = workdir / "hyperlink-output.docx"
    render_report = workdir / "hyperlink-render.json"
    verify_report = workdir / "hyperlink-verify.json"

    doc = Document()
    doc.add_paragraph("Official website:")
    add_external_hyperlink_paragraph(doc, "Old Site", "https://old.example.com")
    doc.save(template)

    analysis = run_json(
        [
            sys.executable,
            str(SCRIPTS / "analyze_template.py"),
            str(template),
            "--output",
            str(analysis_report),
        ]
    )
    hyperlinks = analysis.get("hyperlinks", [])
    check(hyperlinks, "Analyze did not inventory existing hyperlinks.")
    first_link = hyperlinks[0]
    check(first_link["text"] == "Old Site", "Analyze did not capture hyperlink visible text.")
    check(first_link["target"] == "https://old.example.com", "Analyze did not resolve hyperlink relationship target.")
    check(first_link["spec_candidate"]["locator_type"] == "hyperlink", "Analyze did not propose a hyperlink spec candidate.")

    draft = run_json(
        [
            sys.executable,
            str(SCRIPTS / "draft_spec.py"),
            str(template),
            "--report",
            str(draft_report),
        ]
    )
    check(draft["hyperlink_candidates"], "Draft report did not expose hyperlink candidates.")
    check(not any(field.get("locator_type") == "hyperlink" for field in draft["spec"]["fields"]), "Draft spec should not auto-fill hyperlink candidates.")
    check(any(row.get("type") == "hyperlink_candidate" for row in draft["warnings"]), "Draft report did not warn about manual hyperlink promotion.")

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "official_site",
                        "locator_type": "hyperlink",
                        "locator": {
                            "text": "Old Site",
                            "match": "exact",
                            "target": "https://old.example.com",
                        },
                        "replacement_mode": "hyperlink",
                        "required": True,
                    }
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(
        json.dumps(
            {"fields": {"official_site": {"text": "New Site", "url": "https://new.example.com"}}},
            indent=2,
        ),
        encoding="utf-8",
    )
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(render_report),
        ]
    )
    field = report["fields"][0]
    check(field["status"] == "filled", "Hyperlink field was not filled.")
    check(field["action"] == "hyperlink_replace", "Hyperlink field used the wrong action.")
    check(field["display_text_before"] == "Old Site" and field["display_text"] == "New Site", "Hyperlink display text was not reported correctly.")
    check(field["relationship_updated"], "Hyperlink relationship update was not reported.")
    check(field["target_before"] == "https://old.example.com" and field["target"] == "https://new.example.com", "Hyperlink relationship target report is wrong.")
    check(field["format_check"]["preserved"], "Hyperlink replacement did not preserve run/paragraph formatting.")

    document_xml = docx_xml_text(output, "word/document.xml")
    document_rels_xml = docx_xml_text(output, "word/_rels/document.xml.rels")
    check("<w:hyperlink" in document_xml, "Hyperlink wrapper was not preserved.")
    check("New Site" in document_xml and "Old Site" not in document_xml, "Hyperlink display text was not replaced.")
    check("https://new.example.com" in document_rels_xml and "https://old.example.com" not in document_rels_xml, "Hyperlink relationship target was not updated.")

    verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(verify_report),
            "--render-report",
            str(render_report),
        ]
    )
    check(check_named(verify, "tracked_text_format_preserved")["passed"], "Verifier treated hyperlink text replacement as format drift.")


def test_spec_hyperlink_insert(workdir):
    require_python_docx()
    template = workdir / "hyperlink-insert-template.docx"
    spec = workdir / "hyperlink-insert-spec.json"
    data = workdir / "hyperlink-insert-data.json"
    output = workdir / "hyperlink-insert-output.docx"
    render_report = workdir / "hyperlink-insert-render.json"
    verify_report = workdir / "hyperlink-insert-verify.json"

    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Portal: ")
    link_run = paragraph.add_run("{{portal_link}}")
    link_run.bold = True
    paragraph.add_run(" for onboarding.")
    doc.save(template)

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "portal_link",
                        "locator_type": "placeholder",
                        "locator": {"token": "{{portal_link}}"},
                        "replacement_mode": "insert_hyperlink",
                        "required": True,
                    }
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(
        json.dumps(
            {"fields": {"portal_link": {"text": "Client Portal", "url": "https://portal.example.com/onboarding"}}},
            indent=2,
        ),
        encoding="utf-8",
    )

    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(render_report),
        ]
    )
    field = report["fields"][0]
    check(field["status"] == "filled", "New hyperlink field was not filled.")
    check(field["action"] == "hyperlink_insert", "New hyperlink field used the wrong action.")
    check(field["relationship_created"], "New hyperlink relationship was not created.")
    check(field["target"] == "https://portal.example.com/onboarding", "New hyperlink target was not reported.")
    check(field["format_check"]["expected_structure_change"] == "hyperlink_inserted", "Hyperlink insertion did not mark expected structure change.")

    document_xml = docx_xml_text(output, "word/document.xml")
    document_rels_xml = docx_xml_text(output, "word/_rels/document.xml.rels")
    check("<w:hyperlink" in document_xml and "r:id" in document_xml, "New hyperlink wrapper was not inserted.")
    check("Client Portal" in document_xml and "{{portal_link}}" not in document_xml, "New hyperlink display text was not inserted.")
    check("for onboarding." in document_xml, "Text after hyperlink placeholder was not preserved.")
    check("https://portal.example.com/onboarding" in document_rels_xml, "New hyperlink relationship target is missing.")
    check('TargetMode="External"' in document_rels_xml, "New hyperlink relationship is not external.")

    verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(verify_report),
            "--render-report",
            str(render_report),
        ]
    )
    check(check_named(verify, "tracked_text_format_preserved")["passed"], "Verifier treated expected hyperlink insertion as format drift.")


def test_spec_content_control_text(workdir):
    require_python_docx()
    template = workdir / "content-control-template.docx"
    analysis_report = workdir / "content-control-analysis.json"
    spec = workdir / "content-control-spec.json"
    data = workdir / "content-control-data.json"
    output = workdir / "content-control-output.docx"
    report_path = workdir / "content-control-render.json"
    verify_report = workdir / "content-control-verify.json"

    doc = Document()
    header_paragraph = doc.sections[0].header.paragraphs[0]
    header_paragraph.text = "Header Client Placeholder"
    wrap_paragraph_in_content_control(header_paragraph, tag="header_client", alias="Header Client")
    doc.add_paragraph("Student record")
    name_paragraph = doc.add_paragraph("Placeholder Student Name")
    name_sdt = wrap_paragraph_in_content_control(name_paragraph, tag="student_name", alias="Student Name")
    mark_content_control_showing_placeholder(name_sdt)
    major_paragraph = doc.add_paragraph()
    major_paragraph.add_run("Major: ")
    add_run_content_control(major_paragraph, "Placeholder Major", tag="student_major", alias="Major")
    major_paragraph.add_run(" / verified")
    bio_paragraph = doc.add_paragraph("Placeholder Bio")
    wrap_paragraph_in_content_control(
        bio_paragraph,
        tag="student_bio",
        alias="Student Bio",
        text_control={"multi_line": True},
    )
    doc.save(template)

    analysis = run_json(
        [
            sys.executable,
            str(SCRIPTS / "analyze_template.py"),
            str(template),
            "--output",
            str(analysis_report),
        ]
    )
    controls = {row.get("tag"): row for row in analysis.get("content_controls", [])}
    check({"student_name", "student_major", "student_bio", "header_client"}.issubset(controls), "Content controls were not all discovered.")
    check(controls["student_name"]["part"] == "word/document.xml", "Body content control part was not reported.")
    check(controls["student_major"]["level"] == "inline", "Run-level content control was not reported as inline.")
    check(controls["student_bio"]["kind"] == "text", "Plain text content control kind was not reported.")
    check(controls["student_bio"]["text_control"]["multi_line"] is True, "Plain text content control multiLine was not reported.")
    check(controls["student_bio"]["spec_candidate"]["text_control"]["multi_line"] is True, "Content control spec candidate lost multiLine metadata.")
    check(controls["header_client"]["part"].startswith("word/header"), "Header content control part was not reported.")
    check(controls["header_client"]["spec_candidate"]["part"].startswith("word/header"), "Header content control spec candidate lost its part.")
    check(controls["student_name"]["spec_candidate"]["locator"]["tag"] == "student_name", "Content control spec candidate did not use tag locator.")
    check(controls["student_name"]["showing_placeholder"] is True, "Content control placeholder state was not reported.")
    check(controls["student_name"]["spec_candidate"]["showing_placeholder"] is True, "Content control spec candidate lost placeholder state.")

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "student_name",
                        "locator_type": "content_control",
                        "locator": {"tag": "student_name"},
                        "required": True,
                    },
                    {
                        "key": "student_major",
                        "locator_type": "content_control",
                        "locator": {"alias": "Major"},
                        "required": True,
                    },
                    {
                        "key": "student_bio",
                        "locator_type": "content_control",
                        "locator": {"tag": "student_bio"},
                        "required": True,
                    },
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(
        json.dumps(
            {
                "fields": {
                    "student_name": "Li Xiangrui",
                    "student_major": "AI Workflow Engineering",
                    "student_bio": "Builds AI workflows\nKeeps templates faithful",
                }
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    fields = {row["key"]: row for row in report["fields"]}
    check(fields["student_name"]["status"] == "filled", "Block content control was not filled.")
    check(fields["student_major"]["status"] == "filled", "Run content control was not filled.")
    check(fields["student_bio"]["status"] == "filled", "Plain text multiLine content control was not filled.")
    check(fields["student_name"]["action"] == "content_control_text", "Block content control used the wrong action.")
    check(fields["student_major"]["action"] == "content_control_text", "Run content control used the wrong action.")
    check(fields["student_bio"]["action"] == "content_control_text", "Plain text multiLine content control used the wrong action.")
    check(fields["student_bio"]["text_control"]["multi_line"] is True, "Fill report lost multiLine metadata.")
    check(fields["student_bio"].get("content_control_multiline") is True, "Fill report did not mark the content control as multiLine.")
    check(fields["student_bio"].get("line_breaks_inserted") == 1, "Plain text multiLine content control did not materialize a line break.")
    check(fields["student_bio"]["format_check"].get("expected_structure_change") == "line_breaks_inserted", "Plain text multiLine line break was not marked as expected.")
    check(fields["student_name"].get("placeholder_removed") == 1, "Content control placeholder marker was not removed and reported.")
    check(fields["student_name"].get("expected_structure_change") == "content_control_placeholder_removed", "Content control placeholder removal was not marked as expected.")
    document_xml = docx_xml_text(output, "word/document.xml")
    check("Li Xiangrui" in document_xml and "AI Workflow Engineering" in document_xml and "Keeps templates faithful" in document_xml, "Content control values are missing.")
    check("Placeholder Student Name" not in document_xml and "Placeholder Major" not in document_xml and "Placeholder Bio" not in document_xml, "Content control placeholders remain.")
    check("<w:br" in document_xml, "Plain text multiLine content control did not write a Word line break.")
    check("showingPlcHdr" not in document_xml, "Filled content control still has the placeholder-state marker.")
    check(document_xml.count("<w:sdt") >= 2, "Content control wrapper was not preserved.")
    check("student_name" in document_xml and "student_major" in document_xml, "Content control tag metadata was not preserved.")
    verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(verify_report),
            "--render-report",
            str(report_path),
        ]
    )
    check(check_named(verify, "content_controls_preserved")["passed"], "Verifier did not confirm content controls were preserved.")
    check(check_named(verify, "tracked_text_format_preserved")["passed"], "Verifier did not confirm tracked content-control formatting was preserved.")
    check(verify["tracked_text_format"]["checked_count"] == 3, "Verifier did not check all content-control field formats.")
    check(verify["content_controls"]["template_count"] == verify["content_controls"]["output_count"], "Verifier reported content control count drift.")

    pipeline_output = workdir / "content-control-pipeline-output.docx"
    pipeline_report_dir = workdir / "pipeline-reports"
    pipeline = run_json(
        [
            sys.executable,
            str(SCRIPTS / "render_pipeline.py"),
            "--template",
            str(template),
            "--data",
            str(data),
            "--output",
            str(pipeline_output),
            "--report-dir",
            str(pipeline_report_dir),
            "--auto-content-controls",
        ]
    )
    auto_spec_path = Path(pipeline["reports"]["auto_content_control_spec"])
    check(auto_spec_path.exists(), "Pipeline did not write the auto content-control spec.")
    auto_spec = json.loads(auto_spec_path.read_text(encoding="utf-8"))
    check({field["key"] for field in auto_spec["fields"]} == {"student_name", "student_major", "student_bio"}, "Auto content-control spec matched the wrong fields.")
    auto_fields = {field["key"]: field for field in auto_spec["fields"]}
    check(auto_fields["student_name"]["showing_placeholder"] is True, "Auto content-control spec lost placeholder state.")
    check(auto_fields["student_bio"]["text_control"]["multi_line"] is True, "Auto content-control spec lost multiLine metadata.")
    check(all(row["status"] == "filled" for row in pipeline["render_status"]["fields"]), "Pipeline auto content-control fields were not filled.")
    pipeline_fields = {row["key"]: row for row in pipeline["render_status"]["fields"]}
    check(pipeline_fields["student_name"].get("placeholder_removed") == 1, "Pipeline did not report placeholder marker removal.")
    check(pipeline_fields["student_bio"]["text_control"]["multi_line"] is True, "Pipeline render report lost multiLine metadata.")
    check(pipeline_fields["student_bio"].get("line_breaks_inserted") == 1, "Pipeline did not report the multiLine line break.")
    pipeline_document_xml = docx_xml_text(pipeline_output, "word/document.xml")
    check("Li Xiangrui" in pipeline_document_xml and "Keeps templates faithful" in pipeline_document_xml, "Pipeline auto content-control output is missing filled text.")
    check("<w:br" in pipeline_document_xml, "Pipeline output is missing the multiLine Word line break.")
    check("showingPlcHdr" not in docx_xml_text(pipeline_output, "word/document.xml"), "Pipeline output still has the placeholder-state marker.")
    check(not any(row["name"] == "content_controls_preserved" for row in pipeline["verify_failed_checks"]), "Pipeline verifier failed content control preservation.")
    check(not any(row["name"] == "tracked_text_format_preserved" for row in pipeline["verify_failed_checks"]), "Pipeline verifier failed tracked text format preservation.")


def test_spec_content_control_empty_run_style(workdir):
    require_python_docx()
    template = workdir / "empty-run-content-control-template.docx"
    spec = workdir / "empty-run-content-control-spec.json"
    data = workdir / "empty-run-content-control-data.json"
    output = workdir / "empty-run-content-control-output.docx"
    render_report = workdir / "empty-run-content-control-render.json"
    verify_report = workdir / "empty-run-content-control-verify.json"

    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Styled empty field: ")
    add_empty_styled_run_content_control(paragraph, tag="styled_empty", alias="Styled Empty")
    paragraph.add_run(" / complete")
    doc.save(template)

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "styled_empty",
                        "locator_type": "content_control",
                        "locator": {"tag": "styled_empty"},
                        "required": True,
                    }
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(json.dumps({"fields": {"styled_empty": "Styled Value"}}, indent=2), encoding="utf-8")
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(render_report),
        ]
    )
    field = report["fields"][0]
    check(field["status"] == "filled", "Empty styled run content control was not filled.")
    check(field.get("run_reused") is True, "Empty styled run was not reused for text insertion.")
    check(field["format_check"]["preserved"], "Empty styled run format signature was not preserved.")

    document_xml = docx_xml_text(output, "word/document.xml")
    check("Styled Value" in document_xml, "Styled empty run output is missing the filled value.")
    check("Courier New" in document_xml and "SimSun" in document_xml and 'w:val="28"' in document_xml, "Styled run rPr font/size was not preserved.")
    dom = minidom.parseString(document_xml.encode("utf-8"))
    target_content = None
    for sdt in dom.getElementsByTagName("w:sdt"):
        tag = sdt.getElementsByTagName("w:tag")
        if tag and tag[0].getAttribute("w:val") == "styled_empty":
            contents = sdt.getElementsByTagName("w:sdtContent")
            target_content = contents[0] if contents else None
            break
    check(target_content is not None, "Styled empty content control wrapper was not found in output.")
    direct_paragraphs = [node for node in target_content.childNodes if getattr(node, "tagName", None) == "w:p"]
    direct_runs = [node for node in target_content.childNodes if getattr(node, "tagName", None) == "w:r"]
    check(not direct_paragraphs, "Run-level empty content control was converted into a paragraph.")
    check(len(direct_runs) == 1, "Run-level empty content control did not preserve its single direct run.")

    verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(verify_report),
            "--render-report",
            str(render_report),
        ]
    )
    check(check_named(verify, "tracked_text_format_preserved")["passed"], "Verifier did not confirm empty-run format preservation.")
    check(check_named(verify, "content_controls_preserved")["passed"], "Verifier did not confirm empty-run content control preservation.")


def test_spec_content_control_lock_guard(workdir):
    require_python_docx()
    template = workdir / "locked-content-control-template.docx"
    analysis_report = workdir / "locked-content-control-analysis.json"
    spec = workdir / "locked-content-control-spec.json"
    allowed_spec = workdir / "locked-content-control-allowed-spec.json"
    draft_spec_output = workdir / "locked-content-control-draft-spec.json"
    draft_report = workdir / "locked-content-control-draft-report.json"
    data = workdir / "locked-content-control-data.json"
    blocked_output = workdir / "locked-content-control-blocked-output.docx"
    blocked_report_path = workdir / "locked-content-control-blocked-render.json"
    allowed_output = workdir / "locked-content-control-allowed-output.docx"
    allowed_report_path = workdir / "locked-content-control-allowed-render.json"
    verify_report = workdir / "locked-content-control-verify.json"

    doc = Document()
    paragraph = doc.add_paragraph("Locked value: ")
    add_run_content_control(paragraph, "Placeholder Locked", tag="locked_value", alias="Locked Value", lock="contentLocked")
    doc.save(template)

    analysis = run_json(
        [
            sys.executable,
            str(SCRIPTS / "analyze_template.py"),
            str(template),
            "--output",
            str(analysis_report),
        ]
    )
    controls = {row.get("tag"): row for row in analysis.get("content_controls", [])}
    check(controls["locked_value"]["lock"] == "contentLocked", "Locked content control lock value was not detected.")
    check(controls["locked_value"]["spec_candidate"]["lock"] == "contentLocked", "Locked content control spec candidate lost lock metadata.")

    run_json(
        [
            sys.executable,
            str(SCRIPTS / "draft_spec.py"),
            str(template),
            "--spec-output",
            str(draft_spec_output),
            "--report",
            str(draft_report),
        ]
    )
    draft_spec = json.loads(draft_spec_output.read_text(encoding="utf-8"))
    draft_fields = {row["key"]: row for row in draft_spec["fields"]}
    check(draft_fields["locked_value"]["lock"] == "contentLocked", "Draft spec did not preserve lock metadata.")

    base_spec = {
        "template_source": str(template),
        "fields": [
            {
                "key": "locked_value",
                "locator_type": "content_control",
                "locator": {"tag": "locked_value"},
                "required": True,
            }
        ],
    }
    spec.write_text(json.dumps(base_spec, indent=2), encoding="utf-8")
    data.write_text(json.dumps({"fields": {"locked_value": "Override Locked"}}, indent=2), encoding="utf-8")
    cp = subprocess.run(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(blocked_output),
            "--report",
            str(blocked_report_path),
        ],
        capture_output=True,
        text=True,
        timeout=120,
    )
    check(cp.returncode == 1, "Locked content control should fail by default.")
    blocked = json.loads(cp.stdout)
    blocked_field = blocked["fields"][0]
    check(blocked_field["status"] == "error", "Locked content control did not report an error.")
    check(blocked_field["error"] == "content_control_locked", "Locked content control used the wrong error.")
    check(blocked_field["content_control_lock"]["content_locked"] is True, "Locked content control did not report content_locked.")

    allowed = dict(base_spec)
    allowed["fields"] = [dict(base_spec["fields"][0], allow_locked_content_control=True)]
    allowed_spec.write_text(json.dumps(allowed, indent=2), encoding="utf-8")
    allowed_report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(allowed_spec),
            "--data",
            str(data),
            "--output",
            str(allowed_output),
            "--report",
            str(allowed_report_path),
        ]
    )
    allowed_field = allowed_report["fields"][0]
    check(allowed_field["status"] == "filled", "Locked content control override did not fill.")
    check(allowed_field["content_control_lock"]["lock"] == "contentLocked", "Locked content control override did not report lock metadata.")
    document_xml = docx_xml_text(allowed_output, "word/document.xml")
    check("Override Locked" in document_xml, "Locked content control override output is missing text.")
    check("w:lock" in document_xml and "contentLocked" in document_xml, "Locked content control metadata was not preserved.")

    verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(allowed_output),
            "--report",
            str(verify_report),
            "--render-report",
            str(allowed_report_path),
        ]
    )
    check(check_named(verify, "content_controls_preserved")["passed"], "Verifier did not confirm locked content-control preservation.")


def test_spec_content_control_databinding(workdir):
    require_python_docx()
    template = workdir / "databinding-template.docx"
    analysis_report = workdir / "databinding-analysis.json"
    spec = workdir / "databinding-spec.json"
    draft_spec_output = workdir / "databinding-draft-spec.json"
    draft_report = workdir / "databinding-draft-report.json"
    data = workdir / "databinding-data.json"
    output = workdir / "databinding-output.docx"
    report_path = workdir / "databinding-render.json"
    verify_report = workdir / "databinding-verify.json"

    store_item_id = "{11111111-2222-3333-4444-555555555555}"
    xpath = "/contract/client_name[1]"
    doc = Document()
    paragraph = doc.add_paragraph("Client: ")
    add_databound_run_content_control(paragraph, "Placeholder Client", xpath=xpath, store_item_id=store_item_id)
    doc.save(template)
    patch_docx_with_custom_xml(template, store_item_id, "<contract><client_name>Placeholder Client</client_name></contract>")

    analysis = run_json(
        [
            sys.executable,
            str(SCRIPTS / "analyze_template.py"),
            str(template),
            "--output",
            str(analysis_report),
        ]
    )
    binding_control = next(row for row in analysis.get("content_controls", []) if (row.get("binding") or {}).get("xpath") == xpath)
    check(binding_control.get("tag") is None and binding_control.get("alias") is None, "Binding-only fixture unexpectedly has tag or alias.")
    check(binding_control["binding"]["store_item_id"] == store_item_id, "dataBinding storeItemID was not extracted.")
    check(binding_control["spec_candidate"]["key"] == "client_name", "dataBinding spec candidate key was not derived from xpath.")
    check(binding_control["spec_candidate"]["locator"]["binding"]["xpath"] == xpath, "dataBinding spec candidate did not use binding locator.")

    run_json(
        [
            sys.executable,
            str(SCRIPTS / "draft_spec.py"),
            str(template),
            "--spec-output",
            str(draft_spec_output),
            "--report",
            str(draft_report),
        ]
    )
    draft_spec = json.loads(draft_spec_output.read_text(encoding="utf-8"))
    draft_fields = {row["key"]: row for row in draft_spec["fields"]}
    check("client_name" in draft_fields, "Draft spec did not include dataBinding field.")
    check(draft_fields["client_name"]["locator"]["binding"]["xpath"] == xpath, "Draft spec did not preserve dataBinding locator.")

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "client_name",
                        "locator_type": "content_control",
                        "locator": {"binding": {"xpath": xpath, "store_item_id": store_item_id}},
                        "required": True,
                    }
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(json.dumps({"fields": {"client_name": "Acme DataBound"}}, indent=2), encoding="utf-8")
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    field = report["fields"][0]
    check(field["status"] == "filled", "dataBinding content control was not filled.")
    check(field["action"] == "content_control_text", "dataBinding content control used the wrong action.")
    check(field["custom_xml_update"]["updated"] is True, "dataBinding customXml part was not updated.")
    check(field["custom_xml_update"]["part"] == "customXml/item1.xml", "dataBinding customXml update reported the wrong part.")
    document_xml = docx_xml_text(output, "word/document.xml")
    check("Acme DataBound" in document_xml, "dataBinding content-control output is missing filled text.")
    check("Placeholder Client" not in document_xml, "dataBinding content-control placeholder remains.")
    check("w:dataBinding" in document_xml and "client_name" in document_xml, "dataBinding metadata was not preserved.")
    custom_xml = docx_xml_text(output, "customXml/item1.xml")
    check("Acme DataBound" in custom_xml and "Placeholder Client" not in custom_xml, "customXml item value was not updated.")

    verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(verify_report),
            "--render-report",
            str(report_path),
        ]
    )
    check(check_named(verify, "content_controls_preserved")["passed"], "Verifier did not confirm dataBinding content control preservation.")
    check(verify["content_controls"]["template"][0]["binding"]["xpath"] == xpath, "Verifier did not report template dataBinding.")

    pipeline_output = workdir / "databinding-pipeline-output.docx"
    pipeline_report_dir = workdir / "databinding-pipeline-reports"
    pipeline = run_json(
        [
            sys.executable,
            str(SCRIPTS / "render_pipeline.py"),
            "--template",
            str(template),
            "--data",
            str(data),
            "--output",
            str(pipeline_output),
            "--report-dir",
            str(pipeline_report_dir),
            "--auto-content-controls",
        ]
    )
    auto_spec = json.loads(Path(pipeline["reports"]["auto_content_control_spec"]).read_text(encoding="utf-8"))
    check(auto_spec["fields"][0]["key"] == "client_name", "Pipeline auto spec did not derive dataBinding key.")
    check(auto_spec["fields"][0]["locator"]["binding"]["xpath"] == xpath, "Pipeline auto spec did not preserve binding locator.")
    check(pipeline["render_status"]["fields"][0]["action"] == "content_control_text", "Pipeline dataBinding field used the wrong action.")
    check(pipeline["render_status"]["fields"][0]["custom_xml_update"]["updated"] is True, "Pipeline dataBinding did not update customXml.")
    check("Acme DataBound" in docx_xml_text(pipeline_output, "word/document.xml"), "Pipeline dataBinding output is missing filled text.")
    check("Acme DataBound" in docx_xml_text(pipeline_output, "customXml/item1.xml"), "Pipeline customXml output is missing filled text.")
    check(not any(row["name"] == "content_controls_preserved" for row in pipeline["verify_failed_checks"]), "Pipeline verifier failed dataBinding content control preservation.")


def test_spec_content_control_checkbox(workdir):
    require_python_docx()
    template = workdir / "checkbox-template.docx"
    analysis_report = workdir / "checkbox-analysis.json"
    spec = workdir / "checkbox-spec.json"
    draft_spec_output = workdir / "checkbox-draft-spec.json"
    draft_report = workdir / "checkbox-draft-report.json"
    data = workdir / "checkbox-data.json"
    output = workdir / "checkbox-output.docx"
    report_path = workdir / "checkbox-render.json"
    verify_report = workdir / "checkbox-verify.json"

    doc = Document()
    paragraph = doc.add_paragraph("Consent: ")
    add_checkbox_content_control(paragraph, tag="accept_terms", alias="Accept Terms", checked=False)
    paragraph.add_run(" I accept the terms")
    doc.save(template)

    analysis = run_json(
        [
            sys.executable,
            str(SCRIPTS / "analyze_template.py"),
            str(template),
            "--output",
            str(analysis_report),
        ]
    )
    controls = {row.get("tag"): row for row in analysis.get("content_controls", [])}
    check(controls["accept_terms"]["kind"] == "checkBox", "Checkbox content control kind was not detected.")
    check(controls["accept_terms"]["spec_candidate"]["replacement_mode"] == "checkbox", "Checkbox spec candidate did not request checkbox mode.")

    run_json(
        [
            sys.executable,
            str(SCRIPTS / "draft_spec.py"),
            str(template),
            "--spec-output",
            str(draft_spec_output),
            "--report",
            str(draft_report),
        ]
    )
    draft_spec = json.loads(draft_spec_output.read_text(encoding="utf-8"))
    draft_fields = {row["key"]: row for row in draft_spec["fields"]}
    check(draft_fields["accept_terms"]["replacement_mode"] == "checkbox", "Draft spec did not preserve checkbox replacement mode.")

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "accept_terms",
                        "locator_type": "content_control",
                        "locator": {"tag": "accept_terms"},
                        "replacement_mode": "checkbox",
                        "required": True,
                    }
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(json.dumps({"fields": {"accept_terms": True}}, indent=2), encoding="utf-8")
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    field = report["fields"][0]
    check(field["status"] == "filled", "Checkbox content control was not filled.")
    check(field["action"] == "content_control_checkbox", "Checkbox used the wrong action.")
    check(field["checked"] is True, "Checkbox report did not record checked=true.")
    document_xml = docx_xml_text(output, "word/document.xml")
    check("accept_terms" in document_xml, "Checkbox tag metadata was not preserved.")
    check("☒" in document_xml, "Checkbox display text was not updated.")
    check('w14:val="1"' in document_xml or 'w:val="1"' in document_xml, "Checkbox checked state was not set.")

    verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(verify_report),
            "--render-report",
            str(report_path),
        ]
    )
    check(check_named(verify, "content_controls_preserved")["passed"], "Verifier did not confirm checkbox content control preservation.")

    pipeline_output = workdir / "checkbox-pipeline-output.docx"
    pipeline_report_dir = workdir / "checkbox-pipeline-reports"
    pipeline = run_json(
        [
            sys.executable,
            str(SCRIPTS / "render_pipeline.py"),
            "--template",
            str(template),
            "--data",
            str(data),
            "--output",
            str(pipeline_output),
            "--report-dir",
            str(pipeline_report_dir),
            "--auto-content-controls",
        ]
    )
    auto_spec = json.loads(Path(pipeline["reports"]["auto_content_control_spec"]).read_text(encoding="utf-8"))
    check(auto_spec["fields"][0]["replacement_mode"] == "checkbox", "Pipeline auto content-control spec lost checkbox replacement mode.")
    pipeline_field = pipeline["render_status"]["fields"][0]
    check(pipeline_field["action"] == "content_control_checkbox", "Pipeline checkbox used the wrong action.")
    check(pipeline_field["checked"] is True, "Pipeline checkbox did not check the control.")
    check("☒" in docx_xml_text(pipeline_output, "word/document.xml"), "Pipeline checkbox display text was not updated.")
    check(not any(row["name"] == "content_controls_preserved" for row in pipeline["verify_failed_checks"]), "Pipeline verifier failed checkbox content control preservation.")


def test_spec_content_control_choice(workdir):
    require_python_docx()
    template = workdir / "choice-template.docx"
    analysis_report = workdir / "choice-analysis.json"
    spec = workdir / "choice-spec.json"
    draft_spec_output = workdir / "choice-draft-spec.json"
    draft_report = workdir / "choice-draft-report.json"
    data = workdir / "choice-data.json"
    invalid_data = workdir / "choice-invalid-data.json"
    output = workdir / "choice-output.docx"
    report_path = workdir / "choice-render.json"
    invalid_report_path = workdir / "choice-invalid-render.json"
    verify_report = workdir / "choice-verify.json"

    doc = Document()
    paragraph = doc.add_paragraph("Delivery method: ")
    add_choice_content_control(
        paragraph,
        tag="delivery_method",
        alias="Delivery Method",
        kind="dropDownList",
        options=[("Email", "email"), ("Paper", "paper")],
        selected="Email",
    )
    combo_paragraph = doc.add_paragraph("Invoice type: ")
    add_choice_content_control(
        combo_paragraph,
        tag="invoice_type",
        alias="Invoice Type",
        kind="comboBox",
        options=[("VAT Invoice", "vat"), ("Receipt", "receipt")],
        selected="VAT Invoice",
    )
    doc.save(template)

    analysis = run_json(
        [
            sys.executable,
            str(SCRIPTS / "analyze_template.py"),
            str(template),
            "--output",
            str(analysis_report),
        ]
    )
    controls = {row.get("tag"): row for row in analysis.get("content_controls", [])}
    control = controls["delivery_method"]
    combo_control = controls["invoice_type"]
    check(control["kind"] == "dropDownList", "Dropdown content control kind was not detected.")
    check(control["options"] == [{"display_text": "Email", "value": "email"}, {"display_text": "Paper", "value": "paper"}], "Dropdown options were not extracted.")
    check(control["spec_candidate"]["replacement_mode"] == "choice", "Dropdown spec candidate did not request choice mode.")
    check(control["spec_candidate"]["options"][1]["value"] == "paper", "Dropdown spec candidate lost options.")
    check(combo_control["kind"] == "comboBox", "ComboBox content control kind was not detected.")
    check(combo_control["spec_candidate"]["replacement_mode"] == "choice", "ComboBox spec candidate did not request choice mode.")

    run_json(
        [
            sys.executable,
            str(SCRIPTS / "draft_spec.py"),
            str(template),
            "--spec-output",
            str(draft_spec_output),
            "--report",
            str(draft_report),
        ]
    )
    draft_spec = json.loads(draft_spec_output.read_text(encoding="utf-8"))
    draft_fields = {row["key"]: row for row in draft_spec["fields"]}
    check(draft_fields["delivery_method"]["replacement_mode"] == "choice", "Draft spec did not preserve choice replacement mode.")
    check(draft_fields["delivery_method"]["options"][0]["display_text"] == "Email", "Draft spec did not preserve choice options.")
    check(draft_fields["invoice_type"]["replacement_mode"] == "choice", "Draft spec did not preserve comboBox choice mode.")

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "delivery_method",
                        "locator_type": "content_control",
                        "locator": {"tag": "delivery_method"},
                        "replacement_mode": "choice",
                        "required": True,
                    },
                    {
                        "key": "invoice_type",
                        "locator_type": "content_control",
                        "locator": {"tag": "invoice_type"},
                        "replacement_mode": "choice",
                        "required": True,
                    }
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(json.dumps({"fields": {"delivery_method": "paper", "invoice_type": "Special Invoice"}}, indent=2), encoding="utf-8")
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    fields = {row["key"]: row for row in report["fields"]}
    field = fields["delivery_method"]
    combo_field = fields["invoice_type"]
    check(field["status"] == "filled", "Dropdown content control was not filled.")
    check(field["action"] == "content_control_choice", "Dropdown used the wrong action.")
    check(field["selected_display_text"] == "Paper", "Dropdown did not render the matched display text.")
    check(field["selected_value"] == "paper", "Dropdown did not report the matched value.")
    check(field["matched_option"]["value"] == "paper", "Dropdown did not report the matched option.")
    check(combo_field["status"] == "filled", "ComboBox content control was not filled.")
    check(combo_field["action"] == "content_control_choice", "ComboBox used the wrong action.")
    check(combo_field["selected_display_text"] == "Special Invoice", "ComboBox did not render custom display text.")
    check(combo_field["custom_value"] is True, "ComboBox custom value was not reported.")
    document_xml = docx_xml_text(output, "word/document.xml")
    check("delivery_method" in document_xml, "Dropdown tag metadata was not preserved.")
    check("<w:t>Paper</w:t>" in document_xml, "Dropdown display text was not updated.")
    check("<w:t>Special Invoice</w:t>" in document_xml, "ComboBox custom display text was not updated.")
    check("<w:dropDownList>" in document_xml, "Dropdown control definition was not preserved.")
    check("<w:comboBox>" in document_xml, "ComboBox control definition was not preserved.")

    verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(verify_report),
            "--render-report",
            str(report_path),
        ]
    )
    check(check_named(verify, "content_controls_preserved")["passed"], "Verifier did not confirm dropdown content control preservation.")

    invalid_data.write_text(json.dumps({"fields": {"delivery_method": "courier", "invoice_type": "Special Invoice"}}, indent=2), encoding="utf-8")
    cp = subprocess.run(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(invalid_data),
            "--output",
            str(workdir / "choice-invalid-output.docx"),
            "--report",
            str(invalid_report_path),
        ],
        capture_output=True,
        text=True,
        timeout=120,
    )
    check(cp.returncode == 1, "Dropdown invalid option should fail.")
    invalid_report = json.loads(cp.stdout)
    check(invalid_report["fields"][0]["error"] == "choice_value_not_in_options", "Dropdown invalid option reported the wrong error.")

    pipeline_output = workdir / "choice-pipeline-output.docx"
    pipeline_report_dir = workdir / "choice-pipeline-reports"
    pipeline = run_json(
        [
            sys.executable,
            str(SCRIPTS / "render_pipeline.py"),
            "--template",
            str(template),
            "--data",
            str(data),
            "--output",
            str(pipeline_output),
            "--report-dir",
            str(pipeline_report_dir),
            "--auto-content-controls",
        ]
    )
    auto_spec = json.loads(Path(pipeline["reports"]["auto_content_control_spec"]).read_text(encoding="utf-8"))
    auto_fields = {row["key"]: row for row in auto_spec["fields"]}
    check(auto_fields["delivery_method"]["replacement_mode"] == "choice", "Pipeline auto content-control spec lost choice replacement mode.")
    check(auto_fields["delivery_method"]["options"][1]["value"] == "paper", "Pipeline auto content-control spec lost choice options.")
    check(auto_fields["invoice_type"]["replacement_mode"] == "choice", "Pipeline auto content-control spec lost comboBox choice mode.")
    pipeline_fields = {row["key"]: row for row in pipeline["render_status"]["fields"]}
    pipeline_field = pipeline_fields["delivery_method"]
    pipeline_combo_field = pipeline_fields["invoice_type"]
    check(pipeline_field["action"] == "content_control_choice", "Pipeline dropdown used the wrong action.")
    check(pipeline_field["selected_display_text"] == "Paper", "Pipeline dropdown display text was not updated.")
    check(pipeline_combo_field["action"] == "content_control_choice", "Pipeline comboBox used the wrong action.")
    check(pipeline_combo_field["custom_value"] is True, "Pipeline comboBox custom value was not reported.")
    check("Paper" in docx_xml_text(pipeline_output, "word/document.xml"), "Pipeline dropdown output is missing the selected display text.")
    check("Special Invoice" in docx_xml_text(pipeline_output, "word/document.xml"), "Pipeline comboBox output is missing the custom display text.")
    check(not any(row["name"] == "content_controls_preserved" for row in pipeline["verify_failed_checks"]), "Pipeline verifier failed dropdown content control preservation.")


def test_spec_content_control_date(workdir):
    require_python_docx()
    template = workdir / "date-template.docx"
    analysis_report = workdir / "date-analysis.json"
    spec = workdir / "date-spec.json"
    draft_spec_output = workdir / "date-draft-spec.json"
    draft_report = workdir / "date-draft-report.json"
    data = workdir / "date-data.json"
    output = workdir / "date-output.docx"
    report_path = workdir / "date-render.json"
    verify_report = workdir / "date-verify.json"

    doc = Document()
    paragraph = doc.add_paragraph("Contract date: ")
    add_date_content_control(
        paragraph,
        tag="contract_date",
        alias="Contract Date",
        selected="2026-01-01",
        full_date="2026-01-01T00:00:00Z",
        date_format="yyyy-MM-dd",
    )
    doc.save(template)

    analysis = run_json(
        [
            sys.executable,
            str(SCRIPTS / "analyze_template.py"),
            str(template),
            "--output",
            str(analysis_report),
        ]
    )
    controls = {row.get("tag"): row for row in analysis.get("content_controls", [])}
    control = controls["contract_date"]
    check(control["kind"] == "date", "Date content control kind was not detected.")
    check(control["date"]["format"] == "yyyy-MM-dd", "Date format was not extracted.")
    check(control["date"]["full_date"] == "2026-01-01T00:00:00Z", "Date fullDate was not extracted.")
    check(control["spec_candidate"]["replacement_mode"] == "date", "Date spec candidate did not request date mode.")

    run_json(
        [
            sys.executable,
            str(SCRIPTS / "draft_spec.py"),
            str(template),
            "--spec-output",
            str(draft_spec_output),
            "--report",
            str(draft_report),
        ]
    )
    draft_spec = json.loads(draft_spec_output.read_text(encoding="utf-8"))
    draft_fields = {row["key"]: row for row in draft_spec["fields"]}
    check(draft_fields["contract_date"]["replacement_mode"] == "date", "Draft spec did not preserve date replacement mode.")
    check(draft_fields["contract_date"]["date"]["format"] == "yyyy-MM-dd", "Draft spec did not preserve date format.")

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "contract_date",
                        "locator_type": "content_control",
                        "locator": {"tag": "contract_date"},
                        "replacement_mode": "date",
                        "required": True,
                    }
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(json.dumps({"fields": {"contract_date": "2026-07-17"}}, indent=2), encoding="utf-8")
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    field = report["fields"][0]
    check(field["status"] == "filled", "Date content control was not filled.")
    check(field["action"] == "content_control_date", "Date content control used the wrong action.")
    check(field["selected_display_text"] == "2026-07-17", "Date display text was not reported.")
    check(field["full_date"] == "2026-07-17T00:00:00Z", "Date fullDate was not normalized.")
    check(field["date"]["full_date"] == "2026-07-17T00:00:00Z", "Date fullDate was not updated in report.")
    document_xml = docx_xml_text(output, "word/document.xml")
    check("<w:t>2026-07-17</w:t>" in document_xml, "Date display text was not updated.")
    check('w:fullDate="2026-07-17T00:00:00Z"' in document_xml, "Date fullDate attribute was not updated.")
    check(("<w:date " in document_xml or "<w:date>" in document_xml) and "contract_date" in document_xml, "Date control metadata was not preserved.")

    verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(verify_report),
            "--render-report",
            str(report_path),
        ]
    )
    check(check_named(verify, "content_controls_preserved")["passed"], "Verifier did not confirm date content control preservation.")

    pipeline_output = workdir / "date-pipeline-output.docx"
    pipeline_report_dir = workdir / "date-pipeline-reports"
    pipeline = run_json(
        [
            sys.executable,
            str(SCRIPTS / "render_pipeline.py"),
            "--template",
            str(template),
            "--data",
            str(data),
            "--output",
            str(pipeline_output),
            "--report-dir",
            str(pipeline_report_dir),
            "--auto-content-controls",
        ]
    )
    auto_spec = json.loads(Path(pipeline["reports"]["auto_content_control_spec"]).read_text(encoding="utf-8"))
    auto_fields = {row["key"]: row for row in auto_spec["fields"]}
    check(auto_fields["contract_date"]["replacement_mode"] == "date", "Pipeline auto spec lost date replacement mode.")
    check(auto_fields["contract_date"]["date"]["format"] == "yyyy-MM-dd", "Pipeline auto spec lost date properties.")
    pipeline_field = pipeline["render_status"]["fields"][0]
    check(pipeline_field["action"] == "content_control_date", "Pipeline date field used the wrong action.")
    check("2026-07-17" in docx_xml_text(pipeline_output, "word/document.xml"), "Pipeline date output is missing filled text.")
    check('w:fullDate="2026-07-17T00:00:00Z"' in docx_xml_text(pipeline_output, "word/document.xml"), "Pipeline date fullDate was not updated.")
    check(not any(row["name"] == "content_controls_preserved" for row in pipeline["verify_failed_checks"]), "Pipeline verifier failed date content control preservation.")


def test_spec_content_control_repeating_section(workdir):
    require_python_docx()
    template = workdir / "repeating-section-template.docx"
    analysis_report = workdir / "repeating-section-analysis.json"
    spec = workdir / "repeating-section-spec.json"
    draft_spec_output = workdir / "repeating-section-draft-spec.json"
    draft_report = workdir / "repeating-section-draft-report.json"
    data = workdir / "repeating-section-data.json"
    output = workdir / "repeating-section-output.docx"
    report_path = workdir / "repeating-section-render.json"
    verify_report = workdir / "repeating-section-verify.json"

    doc = Document()
    doc.add_paragraph("Line items")
    add_repeating_section_content_control(doc, tag="line_items", alias="Line Items")
    doc.save(template)

    analysis = run_json(
        [
            sys.executable,
            str(SCRIPTS / "analyze_template.py"),
            str(template),
            "--output",
            str(analysis_report),
        ]
    )
    controls = {row.get("tag"): row for row in analysis.get("content_controls", [])}
    check(controls["line_items"]["kind"] == "repeatingSection", "Repeating section was not detected.")
    check(controls["line_items"]["spec_candidate"]["replacement_mode"] == "repeating_section", "Repeating section spec candidate lost replacement mode.")

    run_json(
        [
            sys.executable,
            str(SCRIPTS / "draft_spec.py"),
            str(template),
            "--spec-output",
            str(draft_spec_output),
            "--report",
            str(draft_report),
        ]
    )
    draft_spec = json.loads(draft_spec_output.read_text(encoding="utf-8"))
    draft_fields = {row["key"]: row for row in draft_spec["fields"]}
    check(draft_fields["line_items"]["replacement_mode"] == "repeating_section", "Draft spec did not preserve repeating section mode.")

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "line_items",
                        "locator_type": "content_control",
                        "locator": {"tag": "line_items"},
                        "replacement_mode": "repeating_section",
                        "required": True,
                    }
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(
        json.dumps(
            {
                "fields": {
                    "line_items": [
                        {"name": "Design", "qty": "2"},
                        {"name": "Build", "qty": "3"},
                    ]
                }
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    field = report["fields"][0]
    check(field["status"] == "filled", "Repeating section was not filled.")
    check(field["action"] == "content_control_repeating_section", "Repeating section used the wrong action.")
    check(field["items_created"] == 2, "Repeating section created the wrong number of items.")
    document_xml = docx_xml_text(output, "word/document.xml")
    check("Design x 2" in document_xml and "Build x 3" in document_xml, "Repeating section output is missing cloned item text.")
    check("{{name}}" not in document_xml and "{{qty}}" not in document_xml, "Repeating section placeholders remain.")
    check(document_xml.count("repeatingSectionItem") == 2, "Repeating section item wrapper was not cloned.")

    verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(verify_report),
            "--render-report",
            str(report_path),
        ]
    )
    check(check_named(verify, "content_controls_preserved")["passed"], "Verifier did not allow intentional repeating section clones.")
    check(verify["content_controls"]["allowed_repeating_section_changes"]["enabled"], "Verifier did not enable repeating section allowance from render report.")

    pipeline_output = workdir / "repeating-section-pipeline-output.docx"
    pipeline_report_dir = workdir / "repeating-section-pipeline-reports"
    pipeline = run_json(
        [
            sys.executable,
            str(SCRIPTS / "render_pipeline.py"),
            "--template",
            str(template),
            "--data",
            str(data),
            "--output",
            str(pipeline_output),
            "--report-dir",
            str(pipeline_report_dir),
            "--auto-content-controls",
        ]
    )
    auto_spec = json.loads(Path(pipeline["reports"]["auto_content_control_spec"]).read_text(encoding="utf-8"))
    auto_fields = {row["key"]: row for row in auto_spec["fields"]}
    check(auto_fields["line_items"]["replacement_mode"] == "repeating_section", "Pipeline auto spec lost repeating section mode.")
    check(pipeline["render_status"]["fields"][0]["action"] == "content_control_repeating_section", "Pipeline repeating section used the wrong action.")
    check("Design x 2" in docx_xml_text(pipeline_output, "word/document.xml"), "Pipeline repeating section output is missing cloned item text.")
    check(not any(row["name"] == "content_controls_preserved" for row in pipeline["verify_failed_checks"]), "Pipeline verifier failed repeating section preservation.")


def test_draft_spec_from_template(workdir):
    require_python_docx()
    template = workdir / "draft-template.docx"
    spec_output = workdir / "draft-spec.json"
    data_scaffold = workdir / "draft-data.json"
    draft_report = workdir / "draft-report.json"
    render_data = workdir / "render-data.json"
    output = workdir / "draft-output.docx"
    render_report = workdir / "draft-render.json"
    pipeline_output = workdir / "draft-pipeline-output.docx"
    pipeline_report_dir = workdir / "draft-pipeline-reports"

    doc = Document()
    doc.add_paragraph("Client: {{client_name}}")
    doc.add_paragraph("[此处键入中文标题]")
    doc.add_paragraph("第二章[此处键入一级标题]")
    name_paragraph = doc.add_paragraph("Placeholder Student Name")
    wrap_paragraph_in_content_control(name_paragraph, tag="student_name", alias="Student Name")
    doc.save(template)

    draft = run_json(
        [
            sys.executable,
            str(SCRIPTS / "draft_spec.py"),
            str(template),
            "--spec-output",
            str(spec_output),
            "--data-output",
            str(data_scaffold),
            "--report",
            str(draft_report),
        ]
    )
    spec = json.loads(spec_output.read_text(encoding="utf-8"))
    scaffold = json.loads(data_scaffold.read_text(encoding="utf-8"))
    fields = {row["key"]: row for row in spec["fields"]}

    check("client_name" in fields, "Draft spec missed a placeholder field.")
    check(fields["client_name"]["locator_type"] == "placeholder", "Placeholder field used the wrong locator.")
    check(fields["client_name"]["replacement_mode"] == "token", "Placeholder field did not request token replacement.")
    check("中文标题" in fields, "Draft spec missed a full-paragraph visible sample.")
    check(fields["中文标题"]["locator_type"] == "literal", "Visible sample field used the wrong locator.")
    check("student_name" in fields, "Draft spec missed a tagged content control.")
    check(fields["student_name"]["locator_type"] == "content_control", "Content-control field used the wrong locator.")
    check("[此处键入一级标题]" in scaffold["literal_replacements"], "Draft data scaffold missed an embedded sample literal.")
    check("一级标题" not in fields, "Embedded sample should not be promoted to a spec field without review.")
    check(any(row["type"] == "embedded_visible_sample" for row in draft["warnings"]), "Draft report missed embedded-sample warning.")

    render_data.write_text(
        json.dumps(
            {
                "fields": {
                    "client_name": "Acme Corp",
                    "中文标题": "生成式 AI 学习成本研究",
                    "student_name": "测试作者",
                }
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    render = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec_output),
            "--data",
            str(render_data),
            "--output",
            str(output),
            "--report",
            str(render_report),
        ]
    )
    check(all(row["status"] == "filled" for row in render["fields"]), "Drafted spec fields were not all fillable.")
    render_fields = {row["key"]: row for row in render["fields"]}
    check(render_fields["中文标题"]["action"] == "literal_text_replace", "Drafted literal field did not use structure-preserving literal replacement.")
    check(render_fields["中文标题"]["format_check"]["preserved"], "Drafted literal field did not preserve text formatting.")
    document_xml = docx_xml_text(output, "word/document.xml")
    check("Acme Corp" in document_xml, "Drafted placeholder field was not rendered.")
    check("生成式 AI 学习成本研究" in document_xml, "Drafted visible-sample field was not rendered.")
    check("测试作者" in document_xml, "Drafted content-control field was not rendered.")
    check("student_name" in document_xml and document_xml.count("<w:sdt") >= 1, "Drafted content-control render did not preserve SDT metadata.")

    pipeline = run_json(
        [
            sys.executable,
            str(SCRIPTS / "render_pipeline.py"),
            "--template",
            str(template),
            "--data",
            str(render_data),
            "--output",
            str(pipeline_output),
            "--report-dir",
            str(pipeline_report_dir),
            "--draft-spec",
        ]
    )
    check(pipeline["reports"]["draft_spec"], "Pipeline did not expose the draft spec report path.")
    check(Path(pipeline["reports"]["draft_spec"]).exists(), "Pipeline did not write a draft spec.")
    check(pipeline["draft_status"]["field_count"] >= 3, "Pipeline draft status did not report generated fields.")
    check(all(row["status"] == "filled" for row in pipeline["render_status"]["fields"]), "Pipeline drafted spec fields were not all filled.")
    pipeline_fields = {row["key"]: row for row in pipeline["render_status"]["fields"]}
    check(pipeline_fields["中文标题"]["action"] == "literal_text_replace", "Pipeline drafted literal field did not use structure-preserving literal replacement.")
    check(pipeline_fields["中文标题"]["format_check"]["preserved"], "Pipeline drafted literal field did not preserve text formatting.")
    pipeline_xml = docx_xml_text(pipeline_output, "word/document.xml")
    check("Acme Corp" in pipeline_xml, "Pipeline drafted placeholder field was not rendered.")
    check("生成式 AI 学习成本研究" in pipeline_xml, "Pipeline drafted visible-sample field was not rendered.")
    check("测试作者" in pipeline_xml, "Pipeline drafted content-control field was not rendered.")


def test_request_field_update_pipeline(workdir):
    require_python_docx()
    template = workdir / "field-update-template.docx"
    standalone_output = workdir / "field-update-standalone.docx"
    standalone_report = workdir / "field-update-standalone.json"
    data = workdir / "field-update-data.json"
    pipeline_output = workdir / "field-update-pipeline.docx"
    pipeline_report_dir = workdir / "field-update-reports"
    refresh_output = workdir / "field-refresh-pipeline.docx"
    refresh_report_dir = workdir / "field-refresh-reports"

    doc = Document()
    doc.add_paragraph("Title: {{title}}")
    doc.save(template)
    data.write_text(json.dumps({"fields": {"title": "Pipeline Title"}}, indent=2), encoding="utf-8")

    standalone = run_json(
        [
            sys.executable,
            str(SCRIPTS / "request_field_update.py"),
            str(template),
            "--output",
            str(standalone_output),
            "--report",
            str(standalone_report),
        ]
    )
    check(standalone["status"] == "passed", "Standalone field update request did not pass.")
    check(standalone["changed_parts"] == ["word/settings.xml"], "Standalone field update did not report settings.xml as changed.")
    standalone_settings = docx_xml_text(standalone_output, "word/settings.xml")
    check("updateFields" in standalone_settings and 'val="true"' in standalone_settings, "Standalone DOCX did not request field updates.")

    pipeline = run_json(
        [
            sys.executable,
            str(SCRIPTS / "render_pipeline.py"),
            "--template",
            str(template),
            "--data",
            str(data),
            "--output",
            str(pipeline_output),
            "--report-dir",
            str(pipeline_report_dir),
            "--request-field-update",
        ]
    )
    check(pipeline["reports"]["field_update"], "Pipeline did not expose the field update report path.")
    check(pipeline["reports"].get("summary_markdown"), "Pipeline did not expose the Markdown summary path.")
    pipeline_markdown = Path(pipeline["reports"]["summary_markdown"])
    check(pipeline_markdown.exists(), "Pipeline did not write the Markdown summary.")
    pipeline_markdown_text = pipeline_markdown.read_text(encoding="utf-8")
    check("Template Fidelity Pipeline Report" in pipeline_markdown_text, "Markdown summary is missing its title.")
    check("Field Update Request" in pipeline_markdown_text, "Markdown summary did not report the field update request.")
    check("Delivery Gate" in pipeline_markdown_text, "Markdown summary did not include the delivery gate.")
    check(pipeline["reports"].get("delivery_gate"), "Pipeline did not expose the delivery gate JSON path.")
    check(pipeline["reports"].get("delivery_gate_markdown"), "Pipeline did not expose the delivery gate Markdown path.")
    pipeline_delivery = json.loads(Path(pipeline["reports"]["delivery_gate"]).read_text(encoding="utf-8"))
    check(pipeline_delivery["status"] == "blocked", "Delivery gate should block missing-font delivery by default.")
    check(any(row["code"] == "missing_fonts" for row in pipeline_delivery["blockers"]), "Delivery gate did not block missing fonts.")
    field_update = json.loads(Path(pipeline["reports"]["field_update"]).read_text(encoding="utf-8"))
    verify = json.loads(Path(pipeline["reports"]["verify"]).read_text(encoding="utf-8"))
    check(field_update["status"] == "passed", "Pipeline field update request did not pass.")
    check("word/settings.xml" in field_update["changed_parts"], "Pipeline field update did not change settings.xml.")
    check("word/settings.xml" in verify["allowed_changed_parts"], "Verifier did not allow the reported settings.xml change.")
    check(not any(row["name"] == "invariant_parts_unchanged" for row in pipeline["verify_failed_checks"]), "Pipeline verifier treated the reported settings.xml change as a failure.")
    check("Pipeline Title" in docx_xml_text(pipeline_output, "word/document.xml"), "Pipeline output is missing rendered text.")
    pipeline_settings = docx_xml_text(pipeline_output, "word/settings.xml")
    check("updateFields" in pipeline_settings and 'val="true"' in pipeline_settings, "Pipeline DOCX did not request field updates.")

    refresh_pipeline = run_json(
        [
            sys.executable,
            str(SCRIPTS / "render_pipeline.py"),
            "--template",
            str(template),
            "--data",
            str(data),
            "--output",
            str(refresh_output),
            "--report-dir",
            str(refresh_report_dir),
            "--refresh-fields",
            "--field-refresh-backend",
            "request-only",
        ]
    )
    check(refresh_pipeline["reports"]["field_refresh"], "Pipeline did not expose the field refresh report path.")
    check(refresh_pipeline["reports"]["field_refreshed_output"], "Pipeline did not expose the field-refreshed sidecar DOCX path.")
    check(refresh_pipeline["reports"]["field_refreshed_verify"], "Pipeline did not expose the field-refreshed verification report path.")
    refresh_markdown = Path(refresh_pipeline["reports"]["summary_markdown"])
    check(refresh_markdown.exists(), "Refresh pipeline did not write the Markdown summary.")
    refresh_markdown_text = refresh_markdown.read_text(encoding="utf-8")
    check("Field Refresh Sidecar" in refresh_markdown_text, "Markdown summary did not report field refresh.")
    check("Actual field results refresh: false" in refresh_markdown_text, "Markdown summary overstated request-only field refresh evidence.")
    check("Post-Refresh Package Drift" in refresh_markdown_text, "Markdown summary did not report post-refresh package drift.")
    refresh_delivery = json.loads(Path(refresh_pipeline["reports"]["delivery_gate"]).read_text(encoding="utf-8"))
    check(any(row["code"] == "field_results_not_refreshed" for row in refresh_delivery["warnings"]), "Delivery gate did not warn about request-only field refresh.")
    strict_delivery = subprocess.run(
        [
            sys.executable,
            str(SCRIPTS / "check_delivery.py"),
            str(refresh_pipeline["reports"]["summary"]),
            "--allow-missing-fonts",
            "--require-actual-field-refresh",
        ],
        capture_output=True,
        text=True,
        timeout=120,
    )
    check(strict_delivery.returncode == 2, "Strict delivery gate should block request-only field refresh.")
    strict_delivery_report = json.loads(strict_delivery.stdout)
    check(any(row["code"] == "field_results_not_refreshed" for row in strict_delivery_report["blockers"]), "Strict delivery gate did not block unproven field refresh.")
    blocked_package = subprocess.run(
        [
            sys.executable,
            str(SCRIPTS / "package_delivery.py"),
            str(refresh_pipeline["reports"]["summary"]),
            "--output",
            str(refresh_report_dir / "blocked-delivery.zip"),
        ],
        capture_output=True,
        text=True,
        timeout=120,
    )
    check(blocked_package.returncode == 2, "Package delivery should refuse blocked delivery gates by default.")
    check(not (refresh_report_dir / "blocked-delivery.zip").exists(), "Package delivery wrote a ZIP despite a blocked gate.")
    allowed_package = run_json(
        [
            sys.executable,
            str(SCRIPTS / "package_delivery.py"),
            str(refresh_pipeline["reports"]["summary"]),
            "--output",
            str(refresh_report_dir / "evidence-package.zip"),
            "--manifest-output",
            str(refresh_report_dir / "evidence-package.manifest.json"),
            "--include-blocked",
        ]
    )
    package_path = Path(allowed_package["package"])
    check(package_path.exists(), "Package delivery did not create the evidence ZIP.")
    check(allowed_package["delivery_gate_status"] == "blocked", "Package did not preserve blocked delivery status.")
    manifest = json.loads((refresh_report_dir / "evidence-package.manifest.json").read_text(encoding="utf-8"))
    check(manifest["delivery_gate"]["status"] == "blocked", "Package manifest did not record delivery gate status.")
    with zipfile.ZipFile(package_path) as archive:
        names = set(archive.namelist())
        check("manifest.json" in names and "README.txt" in names, "Package ZIP is missing manifest or README.")
        check(any(name.endswith(".delivery.md") for name in names), "Package ZIP did not include delivery Markdown.")
        check(any(name.endswith(".pipeline.json") for name in names), "Package ZIP did not include pipeline JSON.")
    field_refresh = json.loads(Path(refresh_pipeline["reports"]["field_refresh"]).read_text(encoding="utf-8"))
    field_refreshed_verify = json.loads(Path(refresh_pipeline["reports"]["field_refreshed_verify"]).read_text(encoding="utf-8"))
    check(field_refresh["status"] == "passed", "Request-only field refresh report did not pass.")
    check(field_refresh["selected_backend"] == "request-only", "Request-only field refresh used the wrong backend.")
    check(field_refresh["actual_field_results_refresh"] is False, "Request-only field refresh should not claim actual field-result recalculation.")
    drift = field_refresh.get("post_refresh_drift") or {}
    check(drift.get("status") == "passed", "Request-only field refresh did not report package drift.")
    check(drift.get("changed_part_count") == 0, "Request-only field refresh should not report external-engine package drift.")
    check(not drift.get("core_risk_parts_changed"), "Request-only field refresh should not report core template drift.")
    sidecar_failed_checks = [row["name"] for row in field_refreshed_verify.get("checks", []) if not row.get("passed")]
    check("invariant_parts_unchanged" not in sidecar_failed_checks, "Request-only field-refreshed sidecar changed invariant parts unexpectedly.")
    check("text_part_structure_preserved" not in sidecar_failed_checks, "Request-only field-refreshed sidecar changed text-part structure unexpectedly.")
    check("tracked_text_format_preserved" not in sidecar_failed_checks, "Request-only field-refreshed sidecar changed tracked text formatting unexpectedly.")
    sidecar_settings = docx_xml_text(Path(refresh_pipeline["reports"]["field_refreshed_output"]), "word/settings.xml")
    check("updateFields" in sidecar_settings and 'val="true"' in sidecar_settings, "Field-refreshed sidecar DOCX did not request field updates.")
    check(refresh_pipeline["field_refresh_status"]["actual_field_results_refresh"] is False, "Pipeline summary overstated field refresh evidence.")
    summary_drift = refresh_pipeline["field_refresh_status"]["post_refresh_drift"]
    check(summary_drift["status"] == drift["status"], "Pipeline summary did not report field refresh drift status.")
    check(summary_drift["changed_part_count"] == 0, "Pipeline summary should report no request-only package drift.")
    check(summary_drift["core_risk_parts_changed"] == [], "Pipeline summary should report no request-only core-risk drift.")
    check(summary_drift["text_parts_changed_count"] == 0, "Pipeline summary should report no request-only text-part drift.")
    check(refresh_pipeline["field_refresh_status"]["verify_status"] == field_refreshed_verify["status"], "Pipeline summary did not report sidecar verification status.")
    refresh_module = load_script_module("refresh_fields.py")
    auto_backends = refresh_module.choose_backends("auto")
    check("libreoffice-cli" in auto_backends, "Auto field refresh should try LibreOffice CLI before falling back to request-only.")
    check(auto_backends.index("libreoffice-cli") < auto_backends.index("request-only"), "LibreOffice CLI backend must be tried before request-only fallback.")
    package_drift = refresh_module.compare_docx_packages(refresh_pipeline["reports"]["field_refreshed_output"], refresh_pipeline["reports"]["field_refreshed_output"])
    check(package_drift["status"] == "passed" and package_drift["changed_part_count"] == 0, "Package drift helper should report no changes for identical DOCX paths.")


def test_template_readiness_assessment(workdir):
    require_python_docx()
    l1_template = workdir / "l1-template.docx"
    l1 = Document()
    l1.add_paragraph("论文题目：{{论文标题}}")
    l1.add_paragraph("作者：{{作者}}")
    l1.save(l1_template)

    l1_report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "assess_template_readiness.py"),
            str(l1_template),
            "--output",
            str(workdir / "l1-readiness.json"),
            "--markdown-output",
            str(workdir / "l1-readiness.md"),
        ]
    )
    check(l1_report["level"] == "L1", "Placeholder template should be classified as L1.")
    check(l1_report["signals"]["placeholder_count"] == 2, "Readiness report should count placeholders.")
    check(l1_report["recommended_route"]["mode"] == "render_pipeline", "L1 route should use render_pipeline.")
    check((workdir / "l1-readiness.md").exists(), "Markdown readiness report was not written.")

    complex_template = workdir / "complex-template.docx"
    complex_doc = Document()
    complex_doc.add_paragraph("论文题目：{{论文标题}}")
    add_field_code(complex_doc.add_paragraph("目录："))
    add_office_math(complex_doc.add_paragraph("公式："))
    complex_doc.save(complex_template)

    complex_report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "assess_template_readiness.py"),
            str(complex_template),
            "--output",
            str(workdir / "complex-readiness.json"),
        ]
    )
    complex_signals = complex_report["signals"]["complex_structures"]
    check(complex_signals["field_instruction_count"] >= 1, "Readiness report should count field instructions.")
    check(complex_signals["equation_count"] >= 1, "Readiness report should count equations.")
    check(complex_report["signals"]["complex_risk_count"] >= 2, "Readiness report should expose complex risk count.")
    check(not complex_report["can_claim_100_percent"], "Complex Word structures should block a 100% claim.")
    check(any("Complex Word structures" in item for item in complex_report["blockers"]), "Complex readiness blocker was not reported.")

    complex_analysis = run_json(
        [
            sys.executable,
            str(SCRIPTS / "analyze_template.py"),
            str(complex_template),
            "--output",
            str(workdir / "complex-analysis.json"),
        ]
    )
    check(complex_analysis["complex_structures"]["field_instructions"], "Analysis should expose field instruction samples.")

    sample_template = workdir / "sample-template.docx"
    sample = Document()
    sample.add_paragraph("[此处键入中文标题]")
    sample.add_paragraph("摘要内容。")
    add_bookmark_text(sample.add_paragraph("目  录"), "MTUpdateHome", "", bookmark_id=33)
    sample.add_table(rows=2, cols=2)
    sample.save(sample_template)

    sample_report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "assess_template_readiness.py"),
            str(sample_template),
            "--output",
            str(workdir / "sample-readiness.json"),
        ]
    )
    check(sample_report["level"] == "L2", "Format sample should be classified as L2.")
    check(sample_report["level_name"] == "format-sample-draft-spec", "Plain sample template should use draft-spec route.")
    check(not sample_report["can_claim_100_percent"], "Plain format sample must not allow a 100% claim.")
    check(sample_report["signals"]["visible_sample_count"] >= 2, "Visible sample text should be counted.")
    check(sample_report["signals"]["named_bookmark_count"] == 0, "System bookmarks should not count as stable bookmark fields.")
    check(sample_report["signals"]["system_bookmark_count"] == 1, "System bookmarks should be reported separately.")
    check(sample_report["blockers"], "Plain format sample should explain the missing stable locators.")

    sample_analysis = run_json(
        [
            sys.executable,
            str(SCRIPTS / "analyze_template.py"),
            str(sample_template),
            "--output",
            str(workdir / "sample-analysis.json"),
        ]
    )
    system_bookmarks = [row for row in sample_analysis["bookmarks"] if row.get("system")]
    check(system_bookmarks and system_bookmarks[0]["name"] == "MTUpdateHome", "System bookmark was not identified.")
    check("spec_candidate" not in system_bookmarks[0], "System bookmark should not generate a spec candidate.")

    sample_draft = run_json(
        [
            sys.executable,
            str(SCRIPTS / "draft_spec.py"),
            str(sample_template),
            "--spec-output",
            str(workdir / "sample-draft.spec.json"),
            "--data-output",
            str(workdir / "sample-draft.data.json"),
            "--report",
            str(workdir / "sample-draft.report.json"),
        ]
    )
    check(not sample_draft["bookmark_candidates"], "System bookmark should not be promoted by draft_spec.py.")

    system_only_template = workdir / "system-bookmark-only-template.docx"
    system_only = Document()
    add_bookmark_text(system_only.add_paragraph("目  录"), "MTUpdateHome", "", bookmark_id=34)
    system_only.save(system_only_template)
    system_only_draft = run_json(
        [
            sys.executable,
            str(SCRIPTS / "draft_spec.py"),
            str(system_only_template),
            "--spec-output",
            str(workdir / "system-only-draft.spec.json"),
            "--data-output",
            str(workdir / "system-only-draft.data.json"),
            "--report",
            str(workdir / "system-only-draft.report.json"),
        ]
    )
    check(not system_only_draft["bookmark_candidates"], "System-only bookmark should not be promoted by draft_spec.py.")
    check(any(row["type"] == "no_stable_fields" for row in system_only_draft["warnings"]), "Draft spec should warn when no stable fields are found.")


def test_spec_table_loop(workdir):
    require_python_docx()
    template = workdir / "table-template.docx"
    spec = workdir / "table-spec.json"
    data = workdir / "table-data.json"
    output = workdir / "table-output.docx"
    report_path = workdir / "table-render.json"
    verify_report = workdir / "table-verify.json"

    doc = Document()
    set_split_runs(doc.add_paragraph(), ["Quote for {{client", "_name}}"])
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Qty"
    table.cell(0, 2).text = "Price"
    set_split_runs(table.cell(1, 0).paragraphs[0], ["{{it", "em}}"])
    table.cell(1, 1).text = "{{qty}}"
    table.cell(1, 2).text = "{{price}}"
    doc.save(template)

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "client_name",
                        "locator_type": "placeholder",
                        "locator": {"token": "{{client_name}}"},
                        "replacement_mode": "token",
                        "required": True,
                    }
                ],
                "table_loops": [
                    {
                        "key": "line_items",
                        "locator_type": "table_index",
                        "locator": {"table_index": 0},
                        "row_index": 1,
                        "remove_template_row": True,
                    }
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(
        json.dumps(
            {
                "fields": {"client_name": "Northwind"},
                "tables": {
                    "line_items": [
                        {"item": "Automation", "qty": "1", "price": "12000"},
                        {"item": "Training", "qty": "2", "price": "3000"},
                    ]
                },
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    loop = report["table_loops"][0]
    check(loop["status"] == "filled", "Table loop was not filled.")
    check(loop["rows_inserted"] == 2, "Table loop inserted the wrong number of rows.")
    check(loop["format_check"]["checked"], "Table loop did not emit a format check.")
    check(loop["format_check"]["preserved"], "Table loop row/cell format was not preserved.")
    check(loop["format_check"]["rows_checked"] == 2, "Table loop did not check all inserted rows.")
    document_xml = docx_xml_text(output, "word/document.xml")
    check("Northwind" in document_xml and "Automation" in document_xml and "Training" in document_xml, "Expected table output text is missing.")
    check("{{item}}" not in document_xml and "{{client_name}}" not in document_xml, "Output still has table placeholders.")
    verify = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(verify_report),
            "--render-report",
            str(report_path),
        ]
    )
    check(check_named(verify, "tracked_text_format_preserved")["passed"], "Verifier did not confirm table-loop formatting was preserved.")

    bad_render_report = workdir / "bad-table-format-render.json"
    bad_verify_report = workdir / "bad-table-format-verify.json"
    bad_render_report.write_text(
        json.dumps(
            {
                "table_loops": [
                    {
                        "key": "line_items",
                        "part": "word/document.xml",
                        "format_check": {
                            "checked": True,
                            "scope": "table_loop_rows",
                            "preserved": False,
                            "rows_checked": 1,
                            "failures": [{"row_index": 0, "preserved": False}],
                        },
                    }
                ]
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    bad = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "verify_fidelity.py"),
            str(template),
            str(output),
            "--report",
            str(bad_verify_report),
            "--render-report",
            str(bad_render_report),
        ]
    )
    check(not check_named(bad, "tracked_text_format_preserved")["passed"], "Verifier did not fail a table format drift.")


def test_spec_nested_table_loop(workdir):
    require_python_docx()
    template = workdir / "nested-table-template.docx"
    spec = workdir / "nested-table-spec.json"
    data = workdir / "nested-table-data.json"
    output = workdir / "nested-table-output.docx"
    report_path = workdir / "nested-table-render.json"

    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    outer.cell(0, 0).text = "Outer package cell"
    inner = outer.cell(0, 0).add_table(rows=2, cols=2)
    inner.cell(0, 0).text = "Nested Item"
    inner.cell(0, 1).text = "Nested Qty"
    inner.cell(1, 0).text = "{{item}}"
    inner.cell(1, 1).text = "{{qty}}"
    doc.save(template)

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "table_loops": [
                    {
                        "key": "nested_items",
                        "locator_type": "nested_contains_text",
                        "locator": {"text": "{{item}}"},
                        "row_index": 1,
                        "remove_template_row": True,
                    }
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(
        json.dumps(
            {
                "tables": {
                    "nested_items": [
                        {"item": "Appendix A", "qty": "1"},
                        {"item": "Appendix B", "qty": "2"},
                    ]
                }
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    loop = report["table_loops"][0]
    document_xml = docx_xml_text(output, "word/document.xml")
    check(loop["status"] == "filled", "Nested table loop was not filled.")
    check(loop["rows_inserted"] == 2, "Nested table loop inserted the wrong number of rows.")
    check(loop["table_match"]["table_depth"] == 1, "Nested table loop did not target the inner table.")
    check(loop["table_match"]["tables_matched"] >= 1, "Nested table loop did not report matched tables.")
    check("Outer package cell" in document_xml, "Outer table content was lost.")
    check("Appendix A" in document_xml and "Appendix B" in document_xml, "Nested table output text is missing.")
    check("{{item}}" not in document_xml and "{{qty}}" not in document_xml, "Nested table placeholders were not removed.")


def test_spec_header_footer_part(workdir):
    require_python_docx()
    template = workdir / "part-template.docx"
    spec = workdir / "part-spec.json"
    data = workdir / "part-data.json"
    output = workdir / "part-output.docx"
    report_path = workdir / "part-render.json"

    doc = Document()
    section = doc.sections[0]
    set_split_runs(section.header.paragraphs[0], ["Client: {{client", "_name}}"])
    set_split_runs(section.footer.paragraphs[0], ["Project code: {{project", "_code}}"])
    doc.add_paragraph("Body stays stable.")
    doc.save(template)

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "client_name",
                        "part": "headers",
                        "locator_type": "placeholder",
                        "locator": {"token": "{{client_name}}"},
                        "replacement_mode": "token",
                        "required": True,
                    },
                    {
                        "key": "project_code",
                        "part": "footers",
                        "locator_type": "placeholder",
                        "locator": {"token": "{{project_code}}"},
                        "replacement_mode": "token",
                        "required": True,
                    },
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(
        json.dumps({"fields": {"client_name": "Acme Corp", "project_code": "P-2026"}}, indent=2),
        encoding="utf-8",
    )
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    check(all(row["status"] == "filled" for row in report["fields"]), "Header/footer spec fields were not filled.")
    check({row["part"] for row in report["fields"]} == {"word/header1.xml", "word/footer1.xml"}, "Header/footer fields did not resolve to expected parts.")
    check("Acme Corp" in docx_xml_text(output, "word/header1.xml"), "Header part was not filled.")
    check("P-2026" in docx_xml_text(output, "word/footer1.xml"), "Footer part was not filled.")


def test_spec_section_header_footer_part(workdir):
    require_python_docx()
    template = workdir / "section-part-template.docx"
    spec = workdir / "section-part-spec.json"
    data = workdir / "section-part-data.json"
    output = workdir / "section-part-output.docx"
    report_path = workdir / "section-part-render.json"

    doc = Document()
    first_section = doc.sections[0]
    first_section.header.paragraphs[0].text = "Cover Header Stable"
    first_section.footer.paragraphs[0].text = "Cover Footer Stable"
    doc.add_paragraph("Cover section body.")
    second_section = doc.add_section(WD_SECTION.NEW_PAGE)
    second_section.header.is_linked_to_previous = False
    second_section.footer.is_linked_to_previous = False
    second_section.header.paragraphs[0].text = "Body header: {{body_header}}"
    second_section.footer.paragraphs[0].text = "Body footer: {{body_footer}}"
    doc.add_paragraph("Body section text.")
    doc.save(template)

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "fields": [
                    {
                        "key": "body_header",
                        "part": {"type": "section_header", "section_index": 1, "reference_type": "default"},
                        "locator_type": "placeholder",
                        "locator": {"token": "{{body_header}}"},
                        "replacement_mode": "token",
                        "required": True,
                    },
                    {
                        "key": "body_footer",
                        "part": {"type": "section_footer", "section_index": 1, "reference_type": "default"},
                        "locator_type": "placeholder",
                        "locator": {"token": "{{body_footer}}"},
                        "replacement_mode": "token",
                        "required": True,
                    },
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(
        json.dumps({"fields": {"body_header": "Chapter Header", "body_footer": "Page Footer"}}, indent=2),
        encoding="utf-8",
    )
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    fields = {row["key"]: row for row in report["fields"]}
    check(all(row["status"] == "filled" for row in fields.values()), "Section header/footer fields were not filled.")
    header_resolution = fields["body_header"]["part_resolution"]
    footer_resolution = fields["body_footer"]["part_resolution"]
    check(header_resolution["kind"] == "header" and header_resolution["requested_section_index"] == 1, "Section header part resolution is wrong.")
    check(footer_resolution["kind"] == "footer" and footer_resolution["requested_section_index"] == 1, "Section footer part resolution is wrong.")
    check(not header_resolution["linked_to_previous"] and not footer_resolution["linked_to_previous"], "Section header/footer should not be linked to previous in this fixture.")
    check("Chapter Header" in docx_xml_text(output, fields["body_header"]["part"]), "Section header target part was not filled.")
    check("Page Footer" in docx_xml_text(output, fields["body_footer"]["part"]), "Section footer target part was not filled.")
    check("Cover Header Stable" in docx_xml_text(output, "word/header1.xml"), "Cover header part was unexpectedly changed.")


def test_spec_section_linked_header_guard(workdir):
    require_python_docx()
    template = workdir / "linked-section-template.docx"
    spec = workdir / "linked-section-spec.json"
    allow_spec = workdir / "linked-section-allow-spec.json"
    data = workdir / "linked-section-data.json"
    blocked_output = workdir / "linked-section-output.docx"
    allowed_output = workdir / "linked-section-allow-output.docx"
    blocked_report = workdir / "linked-section-render.json"
    allowed_report = workdir / "linked-section-allow-render.json"

    doc = Document()
    first_section = doc.sections[0]
    first_section.header.paragraphs[0].text = "Shared header: {{body_header}}"
    doc.add_paragraph("Cover section body.")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Body section still linked to previous header.")
    doc.save(template)

    base_spec = {
        "template_source": str(template),
        "fields": [
            {
                "key": "body_header",
                "part": {"type": "section_header", "section_index": 1, "reference_type": "default"},
                "locator_type": "placeholder",
                "locator": {"token": "{{body_header}}"},
                "replacement_mode": "token",
                "required": True,
            }
        ],
    }
    spec.write_text(json.dumps(base_spec, indent=2), encoding="utf-8")
    allowed_spec_data = json.loads(json.dumps(base_spec))
    allowed_spec_data["fields"][0]["allow_linked_section_part"] = True
    allow_spec.write_text(json.dumps(allowed_spec_data, indent=2), encoding="utf-8")
    data.write_text(json.dumps({"fields": {"body_header": "Shared Body Header"}}, indent=2), encoding="utf-8")

    blocked = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(blocked_output),
            "--report",
            str(blocked_report),
        ]
    )
    blocked_field = blocked["fields"][0]
    check(blocked_field["status"] == "error", "Linked section header should be blocked by default.")
    check(blocked_field["error"] == "section_part_linked_to_previous", "Linked section header used the wrong error code.")
    check(blocked_field["part_resolution"]["linked_to_previous"], "Linked section header did not report linked_to_previous.")
    check(blocked_field["part_resolution"]["source_section_index"] == 0, "Linked section header did not resolve to the previous section.")

    allowed = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(allow_spec),
            "--data",
            str(data),
            "--output",
            str(allowed_output),
            "--report",
            str(allowed_report),
        ]
    )
    allowed_field = allowed["fields"][0]
    check(allowed_field["status"] == "filled", "Allowed linked section header was not filled.")
    check(allowed_field["part_resolution"]["linked_to_previous"], "Allowed linked section header did not report linked_to_previous.")
    check("Shared Body Header" in docx_xml_text(allowed_output, allowed_field["part"]), "Allowed linked section header target part was not filled.")


def test_spec_image_field(workdir):
    require_python_docx()
    template = workdir / "image-template.docx"
    spec = workdir / "image-spec.json"
    data = workdir / "image-data.json"
    output = workdir / "image-output.docx"
    report_path = workdir / "image-render.json"
    template_image = workdir / "template-logo.png"
    replacement_image = workdir / "replacement-logo.gif"

    write_png(template_image, (255, 0, 0))
    write_gif(replacement_image)

    doc = Document()
    paragraph = doc.add_paragraph("Logo slot: ")
    paragraph.add_run().add_picture(str(template_image), width=Inches(0.25))
    doc.save(template)

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "image_fields": [
                    {
                        "key": "logo",
                        "locator_type": "image_index",
                        "locator": {"index": 0},
                        "required": True,
                    }
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(json.dumps({"images": {"logo": str(replacement_image)}}, indent=2), encoding="utf-8")

    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    field = report["image_fields"][0]
    media = docx_media_bytes(output)
    content_types_xml = docx_xml_text(output, "[Content_Types].xml")
    check(field["status"] == "filled", "Image field was not filled.")
    check(field["action"] == "image_replace", "Image field did not report image replacement.")
    check(field["relationship_updated"], "Image field did not report relationship update for extension change.")
    check(field["media_path"].endswith(".gif"), "Image field did not update the media target extension.")
    check('Extension="gif"' in content_types_xml and "image/gif" in content_types_xml, "GIF content type was not added.")
    check(replacement_image.read_bytes() in media.values(), "Replacement image bytes are missing from output media.")
    check(template_image.read_bytes() not in media.values(), "Template image bytes were not replaced.")


def test_spec_conditional_blocks(workdir):
    require_python_docx()
    template = workdir / "conditional-template.docx"
    spec = workdir / "conditional-spec.json"
    false_data = workdir / "conditional-false-data.json"
    false_output = workdir / "conditional-false-output.docx"
    false_report = workdir / "conditional-false-render.json"
    true_data = workdir / "conditional-true-data.json"
    true_output = workdir / "conditional-true-output.docx"
    true_report = workdir / "conditional-true-render.json"

    doc = Document()
    doc.add_paragraph("Contract intro.")
    doc.add_paragraph("{{#if warranty}}")
    doc.add_paragraph("Warranty clause: {{warranty_text}}")
    doc.add_paragraph("{{/if warranty}}")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Price"
    table.cell(1, 0).text = "Support"
    table.cell(1, 1).text = "{{support_fee}}"
    doc.save(template)

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "conditional_blocks": [
                    {
                        "key": "warranty_clause",
                        "condition": {"field": "include_warranty", "equals": True},
                        "locator_type": "marker_pair",
                        "locator": {"start": "{{#if warranty}}", "end": "{{/if warranty}}"},
                    },
                    {
                        "key": "support_row",
                        "condition": "include_support",
                        "locator_type": "table_row_contains",
                        "locator": {"text": "{{support_fee}}"},
                    },
                ],
                "fields": [
                    {
                        "key": "warranty_text",
                        "locator_type": "placeholder",
                        "locator": {"token": "{{warranty_text}}"},
                        "replacement_mode": "token",
                    },
                    {
                        "key": "support_fee",
                        "locator_type": "placeholder",
                        "locator": {"token": "{{support_fee}}"},
                        "replacement_mode": "token",
                    },
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    false_data.write_text(
        json.dumps({"fields": {"include_warranty": False, "include_support": False}}, indent=2),
        encoding="utf-8",
    )
    true_data.write_text(
        json.dumps(
            {
                "fields": {
                    "include_warranty": True,
                    "include_support": True,
                    "warranty_text": "One year on-site support",
                    "support_fee": "3000",
                }
            },
            indent=2,
        ),
        encoding="utf-8",
    )

    false_report_json = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(false_data),
            "--output",
            str(false_output),
            "--report",
            str(false_report),
        ]
    )
    false_xml = docx_xml_text(false_output, "word/document.xml")
    check([row["status"] for row in false_report_json["conditional_blocks"]] == ["removed", "removed"], "False conditional blocks were not removed.")
    check("Warranty clause" not in false_xml and "Support" not in false_xml, "False conditional content remains in the output.")
    check("{{#if warranty}}" not in false_xml and "{{support_fee}}" not in false_xml, "False conditional markers/placeholders remain.")

    true_report_json = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(true_data),
            "--output",
            str(true_output),
            "--report",
            str(true_report),
        ]
    )
    true_xml = docx_xml_text(true_output, "word/document.xml")
    check([row["status"] for row in true_report_json["conditional_blocks"]] == ["kept", "kept"], "True conditional blocks were not kept.")
    check("One year on-site support" in true_xml and "3000" in true_xml, "True conditional field values were not rendered.")
    check("{{#if warranty}}" not in true_xml and "{{/if warranty}}" not in true_xml, "True conditional markers were not removed.")


def test_reference_block_clone(workdir):
    require_python_docx()
    template = workdir / "reference-block-template.docx"
    reference = workdir / "reference-block-reference.docx"
    spec = workdir / "reference-block-spec.json"
    data = workdir / "reference-block-data.json"
    output = workdir / "reference-block-output.docx"
    report_path = workdir / "reference-block-render.json"
    pipeline_output = workdir / "reference-block-pipeline-output.docx"
    pipeline_reports = workdir / "reference-block-pipeline-reports"
    template_image = workdir / "template-image.png"
    reference_image = workdir / "reference-image.png"
    write_png(template_image, (255, 0, 0))
    write_png(reference_image, (0, 0, 255))

    template_doc = Document()
    template_doc.add_paragraph("Before body.")
    add_external_hyperlink_paragraph(template_doc, "Existing template hyperlink", "https://example.com/template-link")
    template_doc.add_picture(str(template_image), width=Inches(0.2))
    template_doc.add_paragraph("{{body}}")
    template_doc.add_paragraph("After body.")
    template_doc.save(template)

    reference_doc = Document()
    custom_style = reference_doc.styles.add_style("Thesis Body Custom", WD_STYLE_TYPE.PARAGRAPH)
    custom_style.base_style = reference_doc.styles["Normal"]
    custom_style.font.name = "Courier New"
    numbered_style = reference_doc.styles.add_style("Thesis Numbered Style", WD_STYLE_TYPE.PARAGRAPH)
    numbered_style.base_style = reference_doc.styles["Normal"]
    add_numbering_definition(reference_doc, num_id=77, abstract_id=77, lvl_text_value="StyleRef-%1)")
    set_style_numbering(numbered_style, num_id=77)
    reference_doc.add_paragraph("{{#body-style}}")
    reference_doc.add_paragraph("Reference heading", style="Heading 1")
    reference_doc.add_paragraph("Reference body paragraph", style=custom_style)
    add_custom_numbered_paragraph(reference_doc, "Reference numbered paragraph")
    reference_doc.add_paragraph("Reference style numbered paragraph", style=numbered_style)
    add_external_hyperlink_paragraph(reference_doc, "Reference hyperlink", "https://example.com/reference-link")
    reference_doc.add_picture(str(reference_image), width=Inches(0.2))
    reference_doc.add_paragraph("Reference footnote paragraph")
    reference_doc.add_paragraph("Reference endnote paragraph")
    reference_doc.add_paragraph("Reference comment paragraph")
    reference_doc.add_paragraph("{{/body-style}}")
    reference_doc.save(reference)
    patch_docx_with_note(reference, "Reference footnote paragraph", "Reference footnote content", note_id=5, kind="footnotes")
    patch_docx_with_note(reference, "Reference endnote paragraph", "Reference endnote content", note_id=9, kind="endnotes")
    patch_docx_with_comment(reference, "Reference comment paragraph", "Reference comment content", comment_id=13)

    spec.write_text(
        json.dumps(
            {
                "template_source": str(template),
                "reference_source": str(reference),
                "fields": [
                    {
                        "key": "body",
                        "locator_type": "placeholder",
                        "locator": {"token": "{{body}}"},
                        "replacement_mode": "reference_block",
                        "reference_locator_type": "marker_pair",
                        "reference_locator": {"start": "{{#body-style}}", "end": "{{/body-style}}"},
                        "reference_style_policy": "last",
                        "required": True,
                    }
                ],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    data.write_text(
        json.dumps(
            {
                "fields": {
                    "body": [
                        "生成式 AI 学习成本",
                        "第一段正文保留参考样张的正文样式。",
                        "第二段正文保留直接编号样式。",
                        "第三段正文保留样式内部编号。",
                        "第四段正文保留超链接关系。",
                        "第五段正文保留内嵌图片关系。",
                        "第六段正文保留脚注。",
                        "第七段正文保留尾注。",
                        "第八段正文保留批注。",
                    ]
                }
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    field = report["fields"][0]
    document_xml = docx_xml_text(output, "word/document.xml")
    styles_xml = docx_xml_text(output, "word/styles.xml")
    numbering_xml = docx_xml_text(output, "word/numbering.xml")
    document_rels_xml = docx_xml_text(output, "word/_rels/document.xml.rels")
    footnotes_xml = docx_xml_text(output, "word/footnotes.xml")
    endnotes_xml = docx_xml_text(output, "word/endnotes.xml")
    comments_xml = docx_xml_text(output, "word/comments.xml")
    content_types_xml = docx_xml_text(output, "[Content_Types].xml")
    output_media = docx_media_bytes(output)
    check(field["status"] == "filled", "Reference block field was not filled.")
    check(field["action"] == "replace_reference_block", "Reference block did not replace the placeholder paragraph.")
    check(field["reference_paragraphs_used"] == 9 and field["paragraphs_created"] == 9, "Reference block paragraph counts are wrong.")
    check("ThesisBodyCustom" in field["reference_styles"]["merged_style_ids"], "Reference custom style was not reported as merged.")
    check("ThesisNumberedStyle" in field["reference_styles"]["merged_style_ids"], "Reference numbered style was not reported as merged.")
    check("77" in field["reference_styles"]["merged_style_num_ids"], "Reference style-level numbering was not reported.")
    check(field["reference_numbering"]["imported_num_ids"], "Reference numbering was not reported as imported.")
    check(field["numbering_refs_remapped"] == 1, "Reference numbering references were not remapped.")
    check(field["style_numbering_refs_remapped"] == 1, "Reference style numbering references were not remapped.")
    check(field["reference_relationships"]["imported_relationships"], "Reference hyperlink relationship was not reported as imported.")
    check(field["reference_relationships"]["refs_remapped"] == 1, "Reference hyperlink relationship references were not remapped.")
    imported_hyperlink = field["reference_relationships"]["imported_relationships"][0]
    check(imported_hyperlink["old_relationship_id"] != imported_hyperlink["new_relationship_id"], "Reference hyperlink relationship did not get a target-safe new rId.")
    embedded_images = field["reference_relationships"]["embedded_images"]
    check(embedded_images["imported_relationships"], "Reference embedded image relationship was not reported as imported.")
    check(embedded_images["refs_remapped"] == 1, "Reference embedded image relationship references were not remapped.")
    imported_image = embedded_images["imported_relationships"][0]
    check(imported_image["old_relationship_id"] != imported_image["new_relationship_id"], "Reference image relationship did not get a target-safe new rId.")
    check(imported_image["source_media_path"] != imported_image["media_path"], "Reference image media path did not avoid an existing template media path.")
    check(field["reference_notes"]["footnotes"]["imported_notes"], "Reference footnote was not reported as imported.")
    check(field["reference_notes"]["endnotes"]["imported_notes"], "Reference endnote was not reported as imported.")
    check(field["reference_notes"]["footnotes"]["refs_remapped"] == 1, "Reference footnote references were not remapped.")
    check(field["reference_notes"]["endnotes"]["refs_remapped"] == 1, "Reference endnote references were not remapped.")
    check(field["reference_notes"]["footnotes"]["imported_notes"][0]["old_note_id"] != field["reference_notes"]["footnotes"]["imported_notes"][0]["new_note_id"], "Reference footnote did not get a target-safe new id.")
    check(field["reference_notes"]["endnotes"]["imported_notes"][0]["old_note_id"] != field["reference_notes"]["endnotes"]["imported_notes"][0]["new_note_id"], "Reference endnote did not get a target-safe new id.")
    check(field["reference_comments"]["imported_comments"], "Reference comment was not reported as imported.")
    check(field["reference_comments"]["refs_remapped"] == 3, "Reference comment references were not remapped.")
    check(field["reference_comments"]["imported_comments"][0]["old_comment_id"] != field["reference_comments"]["imported_comments"][0]["new_comment_id"], "Reference comment did not get a target-safe new id.")
    check("生成式 AI 学习成本" in document_xml and "第八段正文保留批注" in document_xml, "Reference block output text is missing.")
    check("{{body}}" not in document_xml and "Reference heading" not in document_xml, "Reference block left placeholder or sample text.")
    check('w:val="Heading1"' in document_xml and 'w:val="ThesisBodyCustom"' in document_xml and 'w:val="ThesisNumberedStyle"' in document_xml, "Reference paragraph styles were not cloned.")
    check('w:styleId="ThesisBodyCustom"' in styles_xml and "Courier New" in styles_xml, "Reference custom style definition was not merged.")
    check('w:styleId="ThesisNumberedStyle"' in styles_xml, "Reference numbered style definition was not merged.")
    check("Ref-%1." in numbering_xml, "Reference numbering definition was not merged.")
    check("StyleRef-%1)" in numbering_xml, "Reference style numbering definition was not merged.")
    check("https://example.com/template-link" in document_rels_xml and "https://example.com/reference-link" in document_rels_xml, "Reference hyperlink relationship target was not merged.")
    check(imported_image["media_path"] in output_media, "Reference image media file was not copied into the output package.")
    check(output_media[imported_image["media_path"]] == reference_image.read_bytes(), "Reference image media bytes were not preserved.")
    check("Reference footnote content" in footnotes_xml and 'w:id="5"' not in footnotes_xml, "Reference footnote content was not copied with a new id.")
    check("Reference endnote content" in endnotes_xml and 'w:id="9"' not in endnotes_xml, "Reference endnote content was not copied with a new id.")
    check("Reference comment content" in comments_xml and 'w:id="13"' not in comments_xml, "Reference comment content was not copied with a new id.")
    check("footnotes.xml" in document_rels_xml and "endnotes.xml" in document_rels_xml and "comments.xml" in document_rels_xml, "Footnote/endnote/comment relationships were not added.")
    check("footnotes+xml" in content_types_xml and "endnotes+xml" in content_types_xml and "comments+xml" in content_types_xml, "Footnote/endnote/comment content types were not added.")
    check('w:val="42"' not in document_xml, "Reference block did not remap old numId from cloned paragraph.")
    check('w:val="77"' not in styles_xml, "Reference block did not remap old numId from cloned style.")
    check("w14:paraId" not in document_xml, "Reference block cloned paragraph tracking IDs.")

    pipeline = run_json(
        [
            sys.executable,
            str(SCRIPTS / "render_pipeline.py"),
            "--spec",
            str(spec),
            "--data",
            str(data),
            "--output",
            str(pipeline_output),
            "--report-dir",
            str(pipeline_reports),
        ]
    )
    verify = json.loads(Path(pipeline["reports"]["verify"]).read_text(encoding="utf-8"))
    invariant = check_named(verify, "invariant_parts_unchanged")
    check(pipeline["status"] in {"passed", "warning"}, "Pipeline failed with merged reference styles.")
    check(invariant["passed"], "Pipeline verifier did not allow intentional style merge.")
    check("word/styles.xml" in verify.get("allowed_changed_parts", []), "Pipeline did not whitelist merged styles.xml.")
    check("word/numbering.xml" in verify.get("allowed_changed_parts", []), "Pipeline did not whitelist merged numbering.xml.")


def test_reference_block_style_conflict_guard(workdir):
    require_python_docx()
    template = workdir / "style-conflict-template.docx"
    reference = workdir / "style-conflict-reference.docx"
    spec = workdir / "style-conflict-spec.json"
    allow_spec = workdir / "style-conflict-allow-spec.json"
    data = workdir / "style-conflict-data.json"
    output = workdir / "style-conflict-output.docx"
    allow_output = workdir / "style-conflict-allow-output.docx"
    report_path = workdir / "style-conflict-render.json"
    allow_report_path = workdir / "style-conflict-allow-render.json"

    template_doc = Document()
    template_style = template_doc.styles.add_style("Shared Body", WD_STYLE_TYPE.PARAGRAPH)
    template_style.base_style = template_doc.styles["Normal"]
    template_style.font.name = "Arial"
    style_id = template_style.style_id
    template_doc.add_paragraph("{{body}}")
    template_doc.save(template)

    reference_doc = Document()
    reference_style = reference_doc.styles.add_style("Shared Body", WD_STYLE_TYPE.PARAGRAPH)
    reference_style.base_style = reference_doc.styles["Normal"]
    reference_style.font.name = "Courier New"
    reference_doc.add_paragraph("{{#body-style}}")
    reference_doc.add_paragraph("Reference styled paragraph", style=reference_style)
    reference_doc.add_paragraph("{{/body-style}}")
    reference_doc.save(reference)

    base_spec = {
        "template_source": str(template),
        "reference_source": str(reference),
        "fields": [
            {
                "key": "body",
                "locator_type": "placeholder",
                "locator": {"token": "{{body}}"},
                "replacement_mode": "reference_block",
                "reference_locator_type": "marker_pair",
                "reference_locator": {"start": "{{#body-style}}", "end": "{{/body-style}}"},
                "required": True,
            }
        ],
    }
    spec.write_text(json.dumps(base_spec, indent=2), encoding="utf-8")
    allowed_spec = json.loads(json.dumps(base_spec))
    allowed_spec["fields"][0]["allow_reference_style_conflicts"] = True
    allow_spec.write_text(json.dumps(allowed_spec, indent=2), encoding="utf-8")
    data.write_text(json.dumps({"fields": {"body": ["Allowed style conflict"]}}, indent=2), encoding="utf-8")

    blocked = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    blocked_field = blocked["fields"][0]
    check(blocked_field["status"] == "error", "Same styleId definition conflict should be blocked by default.")
    check(blocked_field["error"] == "reference_style_conflicts", "Style conflict guard reported the wrong error.")
    blocked_conflicts = blocked_field["reference_styles"]["existing_style_definition_conflicts"]
    check(any(row["style_id"] == style_id for row in blocked_conflicts), "Style definition conflict did not name the shared styleId.")

    allowed = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(allow_spec),
            "--data",
            str(data),
            "--output",
            str(allow_output),
            "--report",
            str(allow_report_path),
        ]
    )
    allowed_field = allowed["fields"][0]
    check(allowed_field["status"] == "filled", "Explicitly allowed style conflict did not render.")
    allowed_conflicts = allowed_field["reference_styles"]["existing_style_definition_conflicts"]
    check(any(row["style_id"] == style_id for row in allowed_conflicts), "Allowed style conflict was not preserved in the report.")
    output_xml = docx_xml_text(allow_output, "word/document.xml")
    check("Allowed style conflict" in output_xml and f'w:val="{style_id}"' in output_xml, "Allowed style conflict output did not keep the referenced style id.")


def test_reference_block_complex_structure_guard(workdir):
    require_python_docx()
    template = workdir / "complex-template.docx"
    reference = workdir / "complex-reference.docx"
    spec = workdir / "complex-spec.json"
    allow_spec = workdir / "complex-allow-spec.json"
    data = workdir / "complex-data.json"
    output = workdir / "complex-output.docx"
    allow_output = workdir / "complex-allow-output.docx"
    report_path = workdir / "complex-render.json"
    allow_report_path = workdir / "complex-allow-render.json"

    template_doc = Document()
    template_doc.add_paragraph("{{body}}")
    template_doc.save(template)

    reference_doc = Document()
    reference_doc.add_paragraph("{{#body-style}}")
    reference_doc.add_paragraph("Reference complex paragraph")
    reference_doc.add_paragraph("{{/body-style}}")
    reference_doc.save(reference)
    patch_docx_with_complex_reference_structures(reference, "Reference complex paragraph")

    base_spec = {
        "template_source": str(template),
        "reference_source": str(reference),
        "fields": [
            {
                "key": "body",
                "locator_type": "placeholder",
                "locator": {"token": "{{body}}"},
                "replacement_mode": "reference_block",
                "reference_locator_type": "marker_pair",
                "reference_locator": {"start": "{{#body-style}}", "end": "{{/body-style}}"},
                "required": True,
            }
        ],
    }
    spec.write_text(json.dumps(base_spec, indent=2), encoding="utf-8")
    allowed_spec = json.loads(json.dumps(base_spec))
    allowed_spec["fields"][0]["allow_reference_complex_structures"] = True
    allow_spec.write_text(json.dumps(allowed_spec, indent=2), encoding="utf-8")
    data.write_text(json.dumps({"fields": {"body": ["Allowed complex output"]}}, indent=2), encoding="utf-8")

    blocked = run_allow_warning(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(spec),
            "--data",
            str(data),
            "--output",
            str(output),
            "--report",
            str(report_path),
        ]
    )
    blocked_field = blocked["fields"][0]
    blocked_summary = blocked_field["reference_complex_structures"]
    check(blocked_field["status"] == "error", "Complex reference block should be blocked by default.")
    check(blocked_field["error"] == "reference_complex_structures_present", "Complex reference block used the wrong error code.")
    check(blocked_summary["field_codes"]["by_tag"].get("fldSimple") == 1, "Reference field code was not reported.")
    check(blocked_summary["tracked_revisions"]["by_tag"].get("ins") == 1, "Reference tracked revision was not reported.")

    allowed = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(allow_spec),
            "--data",
            str(data),
            "--output",
            str(allow_output),
            "--report",
            str(allow_report_path),
        ]
    )
    allowed_field = allowed["fields"][0]
    allowed_summary = allowed_field["reference_complex_structures"]
    allow_document_xml = docx_xml_text(allow_output, "word/document.xml")
    check(allowed_field["status"] == "filled", "Allowed complex reference block should still render.")
    check(allowed_summary["has_complex_structures"], "Allowed complex reference block did not report complex structures.")
    check("w:fldSimple" in allow_document_xml and "w:ins" in allow_document_xml, "Allowed complex structures were not preserved in output XML.")


def test_infer_template_smoke(workdir):
    require_python_docx()
    sample_a = workdir / "sample-a.docx"
    sample_b = workdir / "sample-b.docx"
    inferred_template = workdir / "inferred-template.docx"
    inferred_spec = workdir / "inferred-spec.json"
    report_path = workdir / "infer.json"

    for path, client, rows in [
        (sample_a, "Acme", [("Automation", "1"), ("Training", "2")]),
        (sample_b, "Globex", [("Consulting", "3")]),
    ]:
        doc = Document()
        doc.sections[0].header.paragraphs[0].text = f"Header Client: {client}"
        doc.add_paragraph(f"Body Client: {client}")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Qty"
        for item, qty in rows:
            cells = table.add_row().cells
            cells[0].text = item
            cells[1].text = qty
        doc.save(path)

    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "infer_template.py"),
            str(sample_a),
            str(sample_b),
            "--report",
            str(report_path),
            "--spec-output",
            str(inferred_spec),
            "--template-output",
            str(inferred_template),
        ]
    )
    check(report["paragraph_fields"], "Template inference did not find paragraph fields.")
    check(report["text_part_fields"], "Template inference did not find header/footer text-part fields.")
    check(report["table_loops"], "Template inference did not find table loops.")
    check(inferred_template.exists(), "Inferred placeholder template was not written.")
    check(inferred_spec.exists(), "Inferred spec was not written.")
    check("{{" in docx_xml_text(inferred_template, "word/header1.xml"), "Inferred template did not write a header placeholder.")

    render_data = workdir / "inferred-render-data.json"
    rendered = workdir / "inferred-rendered.docx"
    render_report = workdir / "inferred-render.json"
    field_values = {field["key"]: "Initech" for field in report["fields"]}
    table_loop = report["table_loops"][0]
    table_rows = [
        {column["key"]: f"value-{idx}-{column['cell_index']}" for column in table_loop["columns"]}
        for idx in range(2)
    ]
    render_data.write_text(
        json.dumps({"fields": field_values, "tables": {table_loop["key"]: table_rows}}, indent=2),
        encoding="utf-8",
    )
    render_report_json = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(inferred_spec),
            "--data",
            str(render_data),
            "--output",
            str(rendered),
            "--report",
            str(render_report),
        ]
    )
    check(all(row["status"] == "filled" for row in render_report_json["fields"]), "Inferred fields were not filled.")
    check("Initech" in docx_xml_text(rendered, "word/header1.xml"), "Inferred header field was not rendered.")


def test_infer_template_conditional_paragraphs(workdir):
    require_python_docx()
    sample_a = workdir / "conditional-sample-a.docx"
    sample_b = workdir / "conditional-sample-b.docx"
    inferred_template = workdir / "conditional-inferred-template.docx"
    inferred_spec = workdir / "conditional-inferred-spec.json"
    report_path = workdir / "conditional-infer.json"

    doc_a = Document()
    doc_a.add_paragraph("Contract intro")
    doc_a.add_paragraph("Warranty: one year on-site support")
    doc_a.add_paragraph("Contract end")
    doc_a.save(sample_a)

    doc_b = Document()
    doc_b.add_paragraph("Contract intro")
    doc_b.add_paragraph("Contract end")
    doc_b.save(sample_b)

    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "infer_template.py"),
            str(sample_a),
            str(sample_b),
            "--report",
            str(report_path),
            "--spec-output",
            str(inferred_spec),
            "--template-output",
            str(inferred_template),
        ]
    )
    check(report["conditional_blocks"], "Template inference did not find conditional paragraphs.")
    block = report["conditional_blocks"][0]
    key = block["key"]
    check(block["paragraph_indices"] == [1], "Conditional paragraph index was wrong.")
    check(block["spec"]["locator_type"] == "paragraph_contains", "Non-template conditional spec should use paragraph_contains.")
    check(report["spec"]["conditional_blocks"][0]["locator_type"] == "marker_pair", "Placeholder template spec should use marker_pair.")
    template_xml = docx_xml_text(inferred_template, "word/document.xml")
    check("{{#if " + key + "}}" in template_xml and "{{/if " + key + "}}" in template_xml, "Inferred template did not insert conditional markers.")

    false_data = workdir / "conditional-false-data.json"
    false_output = workdir / "conditional-false-output.docx"
    false_report = workdir / "conditional-false-render.json"
    true_data = workdir / "conditional-true-data.json"
    true_output = workdir / "conditional-true-output.docx"
    true_report = workdir / "conditional-true-render.json"
    false_data.write_text(json.dumps({"conditions": {key: False}}, indent=2), encoding="utf-8")
    true_data.write_text(json.dumps({"conditions": {key: True}}, indent=2), encoding="utf-8")

    false_render = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(inferred_spec),
            "--data",
            str(false_data),
            "--output",
            str(false_output),
            "--report",
            str(false_report),
        ]
    )
    false_xml = docx_xml_text(false_output, "word/document.xml")
    check(false_render["conditional_blocks"][0]["status"] == "removed", "False conditional paragraph was not removed.")
    check("Warranty: one year on-site support" not in false_xml, "False conditional paragraph remained in output.")
    check("{{#if " not in false_xml and "{{/if " not in false_xml, "False conditional markers remained in output.")

    true_render = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(inferred_spec),
            "--data",
            str(true_data),
            "--output",
            str(true_output),
            "--report",
            str(true_report),
        ]
    )
    true_xml = docx_xml_text(true_output, "word/document.xml")
    check(true_render["conditional_blocks"][0]["status"] == "kept", "True conditional paragraph was not kept.")
    check("Warranty: one year on-site support" in true_xml, "True conditional paragraph was removed.")
    check("{{#if " not in true_xml and "{{/if " not in true_xml, "True conditional markers remained in output.")


def test_infer_template_conditional_table_rows(workdir):
    require_python_docx()
    sample_a = workdir / "conditional-row-sample-a.docx"
    sample_b = workdir / "conditional-row-sample-b.docx"
    inferred_template = workdir / "conditional-row-inferred-template.docx"
    inferred_spec = workdir / "conditional-row-inferred-spec.json"
    report_path = workdir / "conditional-row-infer.json"

    for path, include_warranty in [(sample_a, True), (sample_b, False)]:
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Amount"
        row = table.add_row().cells
        row[0].text = "Base package"
        row[1].text = "10000"
        if include_warranty:
            row = table.add_row().cells
            row[0].text = "Warranty support"
            row[1].text = "3000"
        row = table.add_row().cells
        row[0].text = "Total"
        row[1].text = "13000" if include_warranty else "10000"
        doc.save(path)

    report = run_json(
        [
            sys.executable,
            str(SCRIPTS / "infer_template.py"),
            str(sample_a),
            str(sample_b),
            "--report",
            str(report_path),
            "--spec-output",
            str(inferred_spec),
            "--template-output",
            str(inferred_template),
        ]
    )
    table_row_blocks = [row for row in report["conditional_blocks"] if row.get("kind") == "conditional_table_row"]
    check(table_row_blocks, "Template inference did not find conditional table rows.")
    block = table_row_blocks[0]
    key = block["key"]
    check(block["table_index"] == 0 and block["row_index"] == 2, "Conditional table row location was wrong.")
    check(block["anchor_text"] == "Warranty support", "Conditional table row anchor was wrong.")
    check(report["spec"]["conditional_blocks"][0]["locator_type"] == "table_row_contains", "Conditional table row spec used the wrong locator.")
    check(report["spec"]["conditional_blocks"][0]["locator"]["table_index"] == 0, "Conditional table row spec did not pin the table index.")
    check(not report["table_loops"], "Optional table row was incorrectly inferred as a table loop.")

    false_data = workdir / "conditional-row-false-data.json"
    false_output = workdir / "conditional-row-false-output.docx"
    false_report = workdir / "conditional-row-false-render.json"
    true_data = workdir / "conditional-row-true-data.json"
    true_output = workdir / "conditional-row-true-output.docx"
    true_report = workdir / "conditional-row-true-render.json"
    false_data.write_text(json.dumps({"conditions": {key: False}}, indent=2), encoding="utf-8")
    true_data.write_text(json.dumps({"conditions": {key: True}}, indent=2), encoding="utf-8")

    false_render = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(inferred_spec),
            "--data",
            str(false_data),
            "--output",
            str(false_output),
            "--report",
            str(false_report),
        ]
    )
    false_xml = docx_xml_text(false_output, "word/document.xml")
    check(false_render["conditional_blocks"][0]["status"] == "removed", "False conditional table row was not removed.")
    check("Warranty support" not in false_xml and ">3000<" not in false_xml, "False conditional table row remained in output.")

    true_render = run_json(
        [
            sys.executable,
            str(SCRIPTS / "fill_by_spec.py"),
            str(inferred_spec),
            "--data",
            str(true_data),
            "--output",
            str(true_output),
            "--report",
            str(true_report),
        ]
    )
    true_xml = docx_xml_text(true_output, "word/document.xml")
    check(true_render["conditional_blocks"][0]["status"] == "kept", "True conditional table row was not kept.")
    check("Warranty support" in true_xml and ">3000<" in true_xml, "True conditional table row was removed.")


def test_visual_ignore_config():
    sys.path.insert(0, str(SCRIPTS))
    from verify_fidelity import load_visual_ignore_config

    config = load_visual_ignore_config(presets=["page-number-footer", "header-footer"])
    check(config["presets"] == ["page-number-footer", "header-footer"], "Preset order was not preserved.")
    check(len(config["regions"]) == 3, "Preset regions were not expanded.")
    check(all(row.get("kind") == "preset" for row in config["regions"]), "Preset regions were not tagged.")
    try:
        load_visual_ignore_config(presets=["missing-preset"])
    except ValueError:
        pass
    else:
        raise AssertionError("Invalid visual ignore preset should be rejected.")


def test_delivery_gate(workdir):
    output = workdir / "ready.docx"
    verify = workdir / "ready.verify.json"
    output.write_text("placeholder", encoding="utf-8")
    verify.write_text("{}", encoding="utf-8")
    gate = load_script_module("check_delivery.py")
    ready_summary = {
        "status": "passed",
        "output": str(output),
        "reports": {"verify": str(verify), "summary": str(workdir / "ready.pipeline.json")},
        "missing_fonts": [],
        "verify_status": "passed",
        "verify_failed_checks": [],
        "render_status": {"unresolved_tokens": []},
        "pdf": {},
        "steps": [{"name": "verify_fidelity", "ok": True, "returncode": 0}],
    }
    ready = gate.evaluate_delivery(ready_summary)
    check(ready["status"] == "ready", "Delivery gate should mark clean summaries as ready.")

    blocked_summary = dict(ready_summary)
    blocked_summary["missing_fonts"] = ["Calibri"]
    blocked_summary["verify_status"] = "warning"
    blocked_summary["verify_failed_checks"] = [{"name": "template_fonts_available", "passed": False, "detail": ["Calibri"]}]
    blocked = gate.evaluate_delivery(blocked_summary)
    check(blocked["status"] == "blocked", "Delivery gate should block missing fonts by default.")
    check(any(row["code"] == "missing_fonts" for row in blocked["blockers"]), "Missing-font blocker was not reported.")

    review = gate.evaluate_delivery(blocked_summary, {"allow_missing_fonts": True})
    check(review["status"] == "needs_review", "Allowed missing fonts should still require review.")
    check(any(row["code"] == "missing_fonts" for row in review["warnings"]), "Allowed missing fonts should be reported as warning.")
    markdown = gate.render_delivery_markdown(blocked)
    check("Delivery Gate Report" in markdown and "missing_fonts" in markdown, "Delivery gate Markdown summary is incomplete.")


def main():
    parser = argparse.ArgumentParser(description="Run deterministic regression tests for template-fidelity-renderer.")
    parser.add_argument("--workdir", help="Directory for generated DOCX fixtures. Defaults to a temporary directory.")
    parser.add_argument("--keep", action="store_true", help="Keep generated fixtures when using a temporary directory.")
    args = parser.parse_args()

    tests = [
        test_header_footer_text_structure,
        test_cross_run_placeholders,
        test_multiline_line_break_rendering,
        test_spec_bookmark_text_replace,
        test_spec_hyperlink_field,
        test_spec_hyperlink_insert,
        test_spec_content_control_text,
        test_spec_content_control_empty_run_style,
        test_spec_content_control_lock_guard,
        test_spec_content_control_databinding,
        test_spec_content_control_checkbox,
        test_spec_content_control_choice,
        test_spec_content_control_date,
        test_spec_content_control_repeating_section,
        test_draft_spec_from_template,
        test_request_field_update_pipeline,
        test_template_readiness_assessment,
        test_spec_table_loop,
        test_spec_nested_table_loop,
        test_spec_header_footer_part,
        test_spec_section_header_footer_part,
        test_spec_section_linked_header_guard,
        test_spec_image_field,
        test_spec_conditional_blocks,
        test_reference_block_clone,
        test_reference_block_style_conflict_guard,
        test_reference_block_complex_structure_guard,
        test_infer_template_smoke,
        test_infer_template_conditional_paragraphs,
        test_infer_template_conditional_table_rows,
        test_delivery_gate,
    ]
    results = []

    if args.workdir:
        workdir = Path(args.workdir).expanduser().resolve()
        workdir.mkdir(parents=True, exist_ok=True)
        cleanup = None
    else:
        cleanup = tempfile.TemporaryDirectory(prefix="template_fidelity_regression_")
        workdir = Path(cleanup.name)

    try:
        for test in tests:
            case_dir = workdir / test.__name__
            case_dir.mkdir(parents=True, exist_ok=True)
            test(case_dir)
            results.append({"name": test.__name__, "passed": True, "workdir": str(case_dir)})
        test_visual_ignore_config()
        results.append({"name": "test_visual_ignore_config", "passed": True, "workdir": None})
        output = {"status": "passed", "workdir": str(workdir), "results": results}
        print(json.dumps(output, ensure_ascii=False, indent=2))
        return 0
    except Exception as exc:
        results.append({"name": "failure", "passed": False, "error": str(exc)})
        output = {"status": "failed", "workdir": str(workdir), "results": results}
        print(json.dumps(output, ensure_ascii=False, indent=2))
        return 1
    finally:
        if cleanup and not args.keep:
            cleanup.cleanup()


if __name__ == "__main__":
    raise SystemExit(main())
